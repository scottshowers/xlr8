"""
Async Upload Router for XLR8
============================

v23 CHANGES (December 2025 - GET HEALTHY Week 2):
- CONSOLIDATED: /upload endpoint moved to smart_router.py
- CONSOLIDATED: /standards/upload endpoint moved to smart_router.py
- This file now contains SUPPORTING endpoints only:
  - /upload/queue-status
  - /upload/status/{job_id}
  - /upload/debug
  - /standards/health
  - /standards/rules
  - /standards/documents
  - /standards/compliance/check/{project_id}
- process_file_background() still exported for smart_router to use

v22 CHANGES:
- Added MetricsService integration for analytics tracking
- Background jobs now record upload metrics (duration, rows, success/fail)
- Metrics available at /api/metrics/summary

v21 CHANGES:
- Added progress_callback support for structured data uploads
- Smooth progress updates during large file processing (no more frozen UI)
- Maps handler progress (5-95%) to job progress (15-65%)

FEATURES:
- Returns immediately after file save (no timeout!)
- Processes in background thread
- Real-time status updates via job polling
- Handles large files gracefully
- Smart Excel parsing (detects blue/colored headers)
- SMART PDF ROUTING - tabular PDFs go to DuckDB!
- CONTENT-BASED ROUTING - routes by content, not file extension
- INTELLIGENCE ANALYSIS - runs on upload for instant insights (Phase 3)

ROUTING LOGIC (Universal Classification Architecture):
- truth_type determines routing, NOT file extension
- REALITY (customer data) → DuckDB (structured queries)
- INTENT (customer docs) → ChromaDB (semantic search)
- REFERENCE (standards) → ChromaDB (semantic search, global)
- CONFIGURATION (mappings) → DuckDB + ChromaDB (both)
- Legacy extension-based routing used when truth_type not specified

Author: XLR8 Team
"""

from fastapi import APIRouter, UploadFile, File, Form, HTTPException, BackgroundTasks, Request, Depends
from typing import Optional
from datetime import datetime
import sys
import os
import json
import threading
import traceback
import hashlib
import PyPDF2
import docx
import pandas as pd
import logging

# Try to import openpyxl for smart Excel parsing
try:
    from openpyxl import load_workbook
    OPENPYXL_AVAILABLE = True
except ImportError:
    OPENPYXL_AVAILABLE = False
    logger_temp = logging.getLogger(__name__)
    logger_temp.warning("openpyxl not available - Excel color detection disabled")

sys.path.insert(0, '/app')
sys.path.insert(0, '/data')

from utils.rag_handler import RAGHandler
from utils.database.models import (
    ProcessingJobModel, DocumentModel, ProjectModel, 
    DocumentRegistryModel, auto_classify_file, LineageModel
)

# Import unified registration service
try:
    from utils.registration_service import RegistrationService, RegistrationSource
    REGISTRATION_SERVICE_AVAILABLE = True
except ImportError:
    try:
        from backend.utils.registration_service import RegistrationService, RegistrationSource
        REGISTRATION_SERVICE_AVAILABLE = True
    except ImportError:
        REGISTRATION_SERVICE_AVAILABLE = False
        logger_temp = logging.getLogger(__name__)
        logger_temp.warning("RegistrationService not available - using legacy registration")

# Import structured data handler for Excel/CSV
try:
    from utils.structured_data_handler import get_structured_handler
    STRUCTURED_HANDLER_AVAILABLE = True
except ImportError:
    STRUCTURED_HANDLER_AVAILABLE = False
    logger_temp = logging.getLogger(__name__)
    logger_temp.warning("Structured data handler not available - Excel/CSV will use RAG only")

# Import metrics service for analytics tracking
try:
    from utils.metrics_service import MetricsService
    METRICS_AVAILABLE = True
except ImportError:
    try:
        from backend.utils.metrics_service import MetricsService
        METRICS_AVAILABLE = True
    except ImportError:
        METRICS_AVAILABLE = False
        logger_temp = logging.getLogger(__name__)
        logger_temp.warning("[UPLOAD] MetricsService not available - metrics will not be recorded")

# Import smart PDF analyzer for tabular PDFs
try:
    from backend.utils.smart_pdf_analyzer import process_pdf_intelligently
    SMART_PDF_AVAILABLE = True
except ImportError:
    SMART_PDF_AVAILABLE = False
    logger_temp = logging.getLogger(__name__)
    logger_temp.warning("Smart PDF analyzer not available - PDFs will use RAG only")

# Import document analyzer for content-based routing
try:
    from backend.utils.document_analyzer import DocumentAnalyzer, DocumentStructure
    DOCUMENT_ANALYZER_AVAILABLE = True
except ImportError:
    DOCUMENT_ANALYZER_AVAILABLE = False
    logger_temp = logging.getLogger(__name__)
    logger_temp.warning("Document analyzer not available - DOCX/TXT will use extension-based routing")

# Import intelligence service for Phase 3 analysis
try:
    from backend.utils.project_intelligence import ProjectIntelligenceService, AnalysisTier
    INTELLIGENCE_AVAILABLE = True
except ImportError:
    INTELLIGENCE_AVAILABLE = False
    logger_temp = logging.getLogger(__name__)
    logger_temp.warning("Project intelligence not available - upload analysis disabled")

logger = logging.getLogger(__name__)

router = APIRouter()


# =============================================================================
# IDEMPOTENT UPLOAD HELPER - Clean existing data before re-upload
# =============================================================================

def _cleanup_existing_file(filename: str, project: str = None, project_id: str = None):
    """
    Remove any existing data for this filename before re-uploading.
    This makes uploads idempotent - same file uploaded twice = replacement, not duplication.
    
    Cleans:
    - ChromaDB chunks
    - DuckDB tables
    - Registry entry
    
    Returns dict with cleanup summary.
    """
    result = {
        'chromadb_chunks_deleted': 0,
        'duckdb_tables_deleted': 0,
        'registry_deleted': False
    }
    
    filename_lower = filename.lower()
    
    # =========================================================================
    # STEP 1: Clean ChromaDB chunks
    # =========================================================================
    try:
        rag = RAGHandler()
        collection = rag.client.get_or_create_collection(name="documents")
        
        # Try multiple metadata fields for matching
        for field in ["filename", "source", "name"]:
            try:
                results = collection.get(where={field: filename})
                if results and results['ids']:
                    collection.delete(ids=results['ids'])
                    result['chromadb_chunks_deleted'] += len(results['ids'])
                    logger.info(f"[CLEANUP] Deleted {len(results['ids'])} ChromaDB chunks (field: {field})")
                    break
            except Exception as ce:
                pass
        
        # Also try case-insensitive match if nothing found
        if result['chromadb_chunks_deleted'] == 0:
            try:
                all_docs = collection.get(include=["metadatas"], limit=10000)
                ids_to_delete = []
                for i, metadata in enumerate(all_docs.get("metadatas", [])):
                    doc_filename = metadata.get("source", metadata.get("filename", ""))
                    if doc_filename and doc_filename.lower() == filename_lower:
                        ids_to_delete.append(all_docs["ids"][i])
                
                if ids_to_delete:
                    collection.delete(ids=ids_to_delete)
                    result['chromadb_chunks_deleted'] = len(ids_to_delete)
                    logger.info(f"[CLEANUP] Deleted {len(ids_to_delete)} ChromaDB chunks (case-insensitive)")
            except Exception as ce2:
                pass
                
    except Exception as chroma_e:
        logger.warning(f"[CLEANUP] ChromaDB cleanup error: {chroma_e}")
    
    # =========================================================================
    # STEP 2: Clean DuckDB tables
    # =========================================================================
    if STRUCTURED_HANDLER_AVAILABLE and project:
        try:
            handler = get_structured_handler()
            conn = handler.conn
            project_lower = project.lower()
            
            # Find tables from _schema_metadata
            try:
                tables_result = conn.execute("""
                    SELECT table_name FROM _schema_metadata 
                    WHERE LOWER(project) = ? AND LOWER(file_name) = ?
                """, [project_lower, filename_lower]).fetchall()
                
                for (table_name,) in tables_result:
                    try:
                        conn.execute(f'DROP TABLE IF EXISTS "{table_name}"')
                        result['duckdb_tables_deleted'] += 1
                        logger.info(f"[CLEANUP] Dropped DuckDB table: {table_name}")
                    except Exception as e:
                        logger.debug(f"Suppressed: {e}")
                
                # Clean metadata
                conn.execute("""
                    DELETE FROM _schema_metadata 
                    WHERE LOWER(project) = ? AND LOWER(file_name) = ?
                """, [project_lower, filename_lower])
                
            except Exception as meta_e:
                pass
            
            # Also check _pdf_tables
            try:
                table_check = conn.execute("""
                    SELECT COUNT(*) FROM information_schema.tables 
                    WHERE table_name = '_pdf_tables'
                """).fetchone()
                
                if table_check and table_check[0] > 0:
                    pdf_result = conn.execute("""
                        SELECT table_name FROM _pdf_tables 
                        WHERE LOWER(source_file) = ?
                    """, [filename_lower]).fetchall()
                    
                    for (table_name,) in pdf_result:
                        try:
                            conn.execute(f'DROP TABLE IF EXISTS "{table_name}"')
                            result['duckdb_tables_deleted'] += 1
                        except Exception as e:
                            logger.debug(f"Suppressed: {e}")
                    
                    conn.execute("""
                        DELETE FROM _pdf_tables WHERE LOWER(source_file) = ?
                    """, [filename_lower])
                    
            except Exception as e:
                logger.debug(f"Suppressed: {e}")
            
            # Clean support tables
            for support_table in ['file_metadata', '_column_profiles']:
                try:
                    conn.execute(f"""
                        DELETE FROM {support_table} 
                        WHERE LOWER(filename) = ? OR LOWER(file_name) = ?
                    """, [filename_lower, filename_lower])
                except Exception as e:
                    logger.debug(f"Suppressed: {e}")
            
            try:
                conn.execute("CHECKPOINT")
            except Exception as e:
                logger.debug(f"Suppressed: {e}")
                
        except Exception as duck_e:
            logger.warning(f"[CLEANUP] DuckDB cleanup error: {duck_e}")
    
    # =========================================================================
    # STEP 3: Unregister from Registry
    # =========================================================================
    try:
        result['registry_deleted'] = DocumentRegistryModel.unregister(filename, project_id)
        if result['registry_deleted']:
            logger.info(f"[CLEANUP] Unregistered {filename} from registry")
    except Exception as reg_e:
        logger.warning(f"[CLEANUP] Registry cleanup error: {reg_e}")
    
    total_cleaned = result['chromadb_chunks_deleted'] + result['duckdb_tables_deleted']
    if total_cleaned > 0 or result['registry_deleted']:
        logger.info(f"[CLEANUP] Cleaned existing data for {filename}: {result}")
    
    return result


# =============================================================================
# DEBUG ENDPOINT - Check what's available
# =============================================================================
@router.get("/upload/debug")
async def debug_features():
    """Debug endpoint to check what features are available"""
    import os
    
    results = {
        "version": "2025-12-15-v21-with-progress-callbacks",  # Update this when deploying
        "smart_pdf_available": SMART_PDF_AVAILABLE,
        "structured_handler_available": STRUCTURED_HANDLER_AVAILABLE,
        "openpyxl_available": OPENPYXL_AVAILABLE,
        "intelligence_available": INTELLIGENCE_AVAILABLE,
        "files_in_utils": [],
        "import_errors": []
    }
    
    # Check what files exist in /app/utils
    try:
        utils_path = "/app/utils"
        if os.path.exists(utils_path):
            results["files_in_utils"] = os.listdir(utils_path)
        else:
            results["files_in_utils"] = ["DIRECTORY NOT FOUND"]
    except Exception as e:
        results["files_in_utils"] = [f"ERROR: {e}"]
    
    # Try imports and capture specific errors
    try:
        from backend.utils.smart_pdf_analyzer import process_pdf_intelligently
        results["smart_pdf_import"] = "SUCCESS"
    except ImportError as e:
        results["smart_pdf_import"] = f"FAILED: {e}"
        results["import_errors"].append(f"smart_pdf_analyzer: {e}")
    except Exception as e:
        results["smart_pdf_import"] = f"ERROR: {e}"
        results["import_errors"].append(f"smart_pdf_analyzer: {e}")
    
    try:
        from utils.structured_data_handler import get_structured_handler
        results["structured_handler_import"] = "SUCCESS"
    except ImportError as e:
        results["structured_handler_import"] = f"FAILED: {e}"
        results["import_errors"].append(f"structured_data_handler: {e}")
    except Exception as e:
        results["structured_handler_import"] = f"ERROR: {e}"
        results["import_errors"].append(f"structured_data_handler: {e}")
    
    try:
        from utils.project_intelligence import ProjectIntelligenceService
        results["intelligence_import"] = "SUCCESS"
    except ImportError as e:
        results["intelligence_import"] = f"FAILED: {e}"
        results["import_errors"].append(f"project_intelligence: {e}")
    except Exception as e:
        results["intelligence_import"] = f"ERROR: {e}"
        results["import_errors"].append(f"project_intelligence: {e}")
    
    # Check pdfplumber
    try:
        import pdfplumber
        results["pdfplumber_available"] = True
    except ImportError:
        results["pdfplumber_available"] = False
        results["import_errors"].append("pdfplumber not installed")
    
    return results


def extract_text(file_path: str) -> str:
    """Extract text from file based on extension"""
    ext = file_path.split('.')[-1].lower()
    
    try:
        if ext == 'pdf':
            # ENHANCED PDF EXTRACTION - try multiple methods
            text = ""
            pages_extracted = 0
            
            # Method 1: Try pdfplumber first (best for tables and structured PDFs)
            try:
                import pdfplumber
                logger.info("[PDF] Trying pdfplumber extraction...")
                with pdfplumber.open(file_path) as pdf:
                    page_texts = []
                    for i, page in enumerate(pdf.pages):
                        page_text = page.extract_text() or ''
                        if page_text.strip():
                            page_texts.append(f"--- Page {i+1} ---\n{page_text}")
                            pages_extracted += 1
                    text = "\n\n".join(page_texts)
                    logger.info(f"[PDF] pdfplumber extracted {pages_extracted} pages, {len(text)} chars")
            except Exception as e:
                logger.warning(f"[PDF] pdfplumber failed: {e}")
            
            # Method 2: If pdfplumber got little/no content, try PyMuPDF
            if len(text) < 500:
                try:
                    import fitz  # PyMuPDF
                    logger.info("[PDF] Trying PyMuPDF extraction...")
                    doc = fitz.open(file_path)
                    page_texts = []
                    for i, page in enumerate(doc):
                        page_text = page.get_text()
                        if page_text.strip():
                            page_texts.append(f"--- Page {i+1} ---\n{page_text}")
                            pages_extracted += 1
                    doc.close()
                    fitz_text = "\n\n".join(page_texts)
                    if len(fitz_text) > len(text):
                        text = fitz_text
                        logger.info(f"[PDF] PyMuPDF extracted {pages_extracted} pages, {len(text)} chars")
                except Exception as e:
                    logger.warning(f"[PDF] PyMuPDF failed: {e}")
            
            # Method 3: Fallback to PyPDF2 if others failed
            if len(text) < 100:
                try:
                    logger.info("[PDF] Trying PyPDF2 extraction...")
                    with open(file_path, 'rb') as f:
                        pdf = PyPDF2.PdfReader(f)
                        page_texts = []
                        for i, page in enumerate(pdf.pages):
                            page_text = page.extract_text() or ''
                            if page_text.strip():
                                page_texts.append(f"--- Page {i+1} ---\n{page_text}")
                                pages_extracted += 1
                        pypdf_text = "\n\n".join(page_texts)
                        if len(pypdf_text) > len(text):
                            text = pypdf_text
                            logger.info(f"[PDF] PyPDF2 extracted {pages_extracted} pages, {len(text)} chars")
                except Exception as e:
                    logger.warning(f"[PDF] PyPDF2 failed: {e}")
            
            # Log final result
            if len(text) < 100:
                logger.error(f"[PDF] All extraction methods failed or returned minimal text ({len(text)} chars)")
            else:
                logger.info(f"[PDF] Final extraction: {len(text)} chars from {pages_extracted} pages")
                
            return text
            
        elif ext == 'docx':
            doc = docx.Document(file_path)
            return '\n'.join([para.text for para in doc.paragraphs])
        elif ext == 'txt':
            with open(file_path, 'r', encoding='utf-8', errors='replace') as f:
                return f.read()
        elif ext == 'md':
            with open(file_path, 'r', encoding='utf-8', errors='replace') as f:
                return f.read()
        elif ext in ['xlsx', 'xls']:
            # For structured data, return a text representation
            df = pd.read_excel(file_path)
            return df.to_string()
        elif ext == 'csv':
            df = pd.read_csv(file_path)
            return df.to_string()
        else:
            # Try to read as text
            with open(file_path, 'r', encoding='utf-8', errors='replace') as f:
                return f.read()
    except Exception as e:
        logger.error(f"Error extracting text from {file_path}: {e}")
        return ""


def detect_excel_header_row(file_path: str) -> int:
    """
    Intelligently detect the header row in an Excel file by looking for:
    - Colored/styled rows (often headers have background colors)
    - Rows with bold text
    - Rows where all cells contain text (not numbers)
    - Common header patterns
    
    Returns the 0-based row index of the likely header row.
    """
    if not OPENPYXL_AVAILABLE:
        return 0  # Default to first row if openpyxl not available
        
    try:
        wb = load_workbook(file_path, data_only=True)
        ws = wb.active
        
        if ws.max_row is None or ws.max_row < 2:
            return 0
            
        # Check first 15 rows for header characteristics
        max_check = min(15, ws.max_row)
        
        for row_idx in range(1, max_check + 1):
            row_cells = list(ws.iter_rows(min_row=row_idx, max_row=row_idx, values_only=False))[0]
            
            # Check for fill color (background color)
            has_fill = False
            has_bold = False
            non_empty_count = 0
            all_text = True
            
            for cell in row_cells:
                if cell.value is not None:
                    non_empty_count += 1
                    
                    # Check if it looks like a number
                    if isinstance(cell.value, (int, float)):
                        all_text = False
                    
                    # Check for fill color
                    if cell.fill and cell.fill.fgColor and cell.fill.fgColor.rgb:
                        rgb = cell.fill.fgColor.rgb
                        # Check if it's not white or transparent
                        if rgb not in ['00000000', 'FFFFFFFF', '00FFFFFF', None]:
                            has_fill = True
                    
                    # Check for bold
                    if cell.font and cell.font.bold:
                        has_bold = True
            
            # If this row has 3+ non-empty cells, has styling, and is mostly text
            if non_empty_count >= 3 and (has_fill or has_bold) and all_text:
                logger.info(f"[EXCEL] Detected header at row {row_idx} (fill={has_fill}, bold={has_bold})")
                wb.close()
                return row_idx - 1  # Convert to 0-based index
        
        wb.close()
        return 0  # Default to first row
        
    except Exception as e:
        logger.warning(f"[EXCEL] Header detection failed: {e}, using row 0")
        return 0


# =============================================================================
# JOB QUEUE FOR SEQUENTIAL PROCESSING
# =============================================================================
import queue
from dataclasses import dataclass, field
from typing import Callable, Any, Tuple

@dataclass(order=True)
class QueuedJob:
    """A job waiting in the queue"""
    priority: int
    job_id: str = field(compare=False)
    func: Callable = field(compare=False)
    args: Tuple = field(compare=False)
    kwargs: dict = field(compare=False, default_factory=dict)
    queued_at: datetime = field(compare=False, default_factory=datetime.now)


class JobQueue:
    """
    Sequential job queue - processes ONE job at a time.
    Prevents Ollama from being overloaded with concurrent requests.
    """
    _instance = None
    _lock = threading.Lock()
    
    def __new__(cls):
        if cls._instance is None:
            with cls._lock:
                if cls._instance is None:
                    cls._instance = super().__new__(cls)
                    cls._instance._initialized = False
        return cls._instance
    
    def __init__(self):
        if self._initialized:
            return
            
        self._queue = queue.PriorityQueue()
        self._current_job = None
        self._processed_count = 0
        self._worker_thread = None
        self._running = True
        self._job_positions = {}  # job_id -> position
        
        # Start worker thread
        self._worker_thread = threading.Thread(target=self._worker_loop, daemon=True)
        self._worker_thread.start()
        
        self._initialized = True
        logger.info("[QUEUE] Job queue initialized with sequential processing")
    
    def enqueue(self, job_id: str, func: Callable, args: Tuple = (), kwargs: dict = None, priority: int = 10) -> dict:
        """Add a job to the queue. Returns immediately with queue position."""
        if kwargs is None:
            kwargs = {}
            
        job = QueuedJob(
            priority=priority,
            job_id=job_id,
            func=func,
            args=args,
            kwargs=kwargs
        )
        
        self._queue.put(job)
        position = self._queue.qsize()
        self._job_positions[job_id] = position
        
        logger.info(f"[QUEUE] Job {job_id} queued at position {position}")
        
        return {
            'queued': True,
            'position': position,
            'queue_size': position,
            'message': f'Queued at position {position}'
        }
    
    def get_position(self, job_id: str) -> int:
        """Get current queue position for a job (0 = processing now)"""
        if self._current_job and self._current_job.job_id == job_id:
            return 0
        return self._job_positions.get(job_id, -1)
    
    def get_status(self) -> dict:
        """Get queue status"""
        return {
            'queue_size': self._queue.qsize(),
            'currently_processing': self._current_job.job_id if self._current_job else None,
            'processed_count': self._processed_count,
            'worker_alive': self._worker_thread.is_alive() if self._worker_thread else False
        }
    
    def _worker_loop(self):
        """Background worker that processes jobs one at a time"""
        logger.info("[QUEUE] Worker thread started")
        
        while self._running:
            try:
                # Wait for a job (timeout allows checking _running flag)
                try:
                    job = self._queue.get(timeout=1.0)
                except queue.Empty:
                    continue
                
                self._current_job = job
                logger.info(f"[QUEUE] Processing job {job.job_id}")
                
                try:
                    # Execute the job function
                    job.func(*job.args, **job.kwargs)
                    self._processed_count += 1
                    logger.info(f"[QUEUE] Job {job.job_id} completed")
                except Exception as e:
                    logger.error(f"[QUEUE] Job {job.job_id} failed: {e}")
                    import traceback
                    logger.error(traceback.format_exc())
                finally:
                    self._current_job = None
                    if job.job_id in self._job_positions:
                        del self._job_positions[job.job_id]
                    self._queue.task_done()
                    self._update_positions()
                    
            except Exception as e:
                logger.error(f"[QUEUE] Worker error: {e}")
    
    def _update_positions(self):
        """Update position tracking after a job completes"""
        # Positions shift down as jobs complete
        for job_id in list(self._job_positions.keys()):
            if self._job_positions[job_id] > 0:
                self._job_positions[job_id] -= 1


# Global queue instance
job_queue = JobQueue()


# =============================================================================
# INTELLIGENCE ANALYSIS HELPER
# =============================================================================

def run_intelligence_analysis(project: str, handler, job_id: str) -> dict:
    """
    Run intelligence analysis on uploaded data.
    
    This is the Phase 3 Universal Analysis Engine integration.
    Runs Tier 1 (instant) and Tier 2 (fast) analysis.
    
    Returns summary of findings, tasks, and alerts.
    """
    if not INTELLIGENCE_AVAILABLE:
        logger.warning("[INTELLIGENCE] Intelligence module not available, skipping analysis")
        return None
    
    try:
        ProcessingJobModel.update_progress(job_id, 72, "Running intelligence analysis...")
        
        intelligence = ProjectIntelligenceService(project, handler)
        analysis = intelligence.analyze(tiers=[AnalysisTier.TIER_1, AnalysisTier.TIER_2])
        
        if 'error' in analysis:
            logger.warning(f"[INTELLIGENCE] Analysis returned error: {analysis.get('error')}")
            return None
        
        # Extract summary
        findings_count = len(analysis.get('findings', []))
        tasks_count = analysis.get('tasks', {}).get('total', 0)
        critical_count = analysis.get('findings_summary', {}).get('critical', 0)
        warning_count = analysis.get('findings_summary', {}).get('warning', 0)
        lookups_count = len(analysis.get('lookups', []))
        relationships_count = len(analysis.get('structure', {}).get('relationships', []))
        
        intelligence_summary = {
            'findings': findings_count,
            'tasks': tasks_count,
            'critical': critical_count,
            'warning': warning_count,
            'lookups_detected': lookups_count,
            'relationships_detected': relationships_count,
            'analyzed_at': analysis.get('analyzed_at'),
            'analysis_time_seconds': analysis.get('analysis_time_seconds')
        }
        
        # Build status message
        if critical_count > 0:
            status_msg = f"⚠️ Intelligence: {critical_count} critical, {warning_count} warnings, {tasks_count} tasks"
        elif findings_count > 0:
            status_msg = f"✅ Intelligence: {findings_count} findings, {tasks_count} tasks"
        else:
            status_msg = f"✅ Intelligence: Data looks clean"
        
        ProcessingJobModel.update_progress(job_id, 78, status_msg)
        logger.info(f"[INTELLIGENCE] Analysis complete: {findings_count} findings, {tasks_count} tasks, {critical_count} critical")
        
        return intelligence_summary
        
    except Exception as e:
        logger.warning(f"[INTELLIGENCE] Analysis failed: {e}")
        import traceback
        logger.warning(traceback.format_exc())
        return None


# =============================================================================
# BACKGROUND PROCESSING
# =============================================================================

def process_file_background(
    job_id: str, 
    file_path: str, 
    filename: str, 
    project: str,
    project_id: Optional[str],
    functional_area: Optional[str],
    file_size: int,
    truth_type: Optional[str] = None,
    content_domain: Optional[str] = None,
    file_hash: Optional[str] = None,
    uploaded_by_id: Optional[str] = None,
    uploaded_by_email: Optional[str] = None
):
    """
    Background processing function - runs in separate thread
    
    Universal Classification Architecture:
    - truth_type determines routing (reality→DuckDB, intent→ChromaDB, etc.)
    - If truth_type not provided, auto-classifies based on filename/extension
    
    Processing:
    1. Classify file (user-provided or auto-detected)
    2. Route based on truth_type:
       - REALITY/CONFIGURATION → DuckDB (structured queries)
       - INTENT/REFERENCE → ChromaDB (semantic search)
    3. Run intelligence analysis (Phase 3)
    4. Update job status throughout
    
    Metadata tracked:
    - file_hash: SHA-256 hash for integrity/dedup
    - uploaded_by_id: User UUID who uploaded
    - uploaded_by_email: User email for quick lookup
    """
    try:
        logger.warning(f"[BACKGROUND] === STARTING JOB {job_id} ===")
        logger.warning(f"[BACKGROUND] filename={filename}, project={project}, project_id={project_id}")
        
        file_ext = filename.split('.')[-1].lower()
        logger.warning(f"[BACKGROUND] Detected file extension: '{file_ext}'")
        
        # Track if DuckDB storage succeeded (for PDFs that go to both DuckDB + ChromaDB)
        duckdb_success = False
        
        # =================================================================
        # TIMING TRACKING - Track processing performance
        # =================================================================
        import time
        timing = {
            'start': time.time(),
            'classification_start': None,
            'classification_end': None,
            'parse_start': None,
            'parse_end': None,
            'embedding_start': None,
            'embedding_end': None,
            'storage_start': None,
            'storage_end': None
        }
        
        # =================================================================
        # CLASSIFICATION - Determine truth_type for routing
        # =================================================================
        timing['classification_start'] = time.time()
        is_global = project.lower() in ['global', '__global__', 'global/universal', 'reference library']
        
        if truth_type:
            # User explicitly provided truth_type
            classification_method = DocumentRegistryModel.CLASS_USER_SELECTED
            classification_confidence = 1.0
            logger.info(f"[BACKGROUND] User-specified truth_type: {truth_type}")
        else:
            # Auto-classify based on filename and extension
            truth_type, classification_method, classification_confidence = auto_classify_file(
                filename=filename,
                file_extension=file_ext,
                project_name=project
            )
            logger.info(f"[BACKGROUND] Auto-classified: {truth_type} (method={classification_method}, confidence={classification_confidence})")
        
        # Determine target storage based on truth_type
        target_storage = DocumentRegistryModel.get_storage_for_truth_type(truth_type)
        logger.info(f"[BACKGROUND] Target storage: {target_storage}")
        
        # Parse content_domain if provided as comma-separated string
        domain_list = []
        if content_domain:
            domain_list = [d.strip() for d in content_domain.split(',') if d.strip()]
        
        timing['classification_end'] = time.time()
        
        # =================================================================
        # Helper to calculate timing in milliseconds
        # =================================================================
        def calc_timing_ms():
            """Calculate all timing values in milliseconds"""
            now = time.time()
            return {
                'processing_time_ms': int((now - timing['start']) * 1000),
                'classification_time_ms': int((timing['classification_end'] - timing['classification_start']) * 1000) if timing['classification_end'] else None,
                'parse_time_ms': int((timing['parse_end'] - timing['parse_start']) * 1000) if timing['parse_end'] and timing['parse_start'] else None,
                'embedding_time_ms': int((timing['embedding_end'] - timing['embedding_start']) * 1000) if timing['embedding_end'] and timing['embedding_start'] else None,
                'storage_time_ms': int((timing['storage_end'] - timing['storage_start']) * 1000) if timing['storage_end'] and timing['storage_start'] else None,
            }
        
        # ROUTE 1: STRUCTURED DATA (Excel/CSV) → DuckDB
        if file_ext in ['xlsx', 'xls', 'csv'] and STRUCTURED_HANDLER_AVAILABLE:
            ProcessingJobModel.update_progress(job_id, 10, "Detected tabular data - storing for SQL queries...")
            timing['parse_start'] = time.time()
            
            try:
                # ===========================================================
                # IDEMPOTENT: Clean existing data for this filename first
                # This prevents duplicate tables if same file uploaded twice
                # ===========================================================
                cleanup_result = _cleanup_existing_file(filename, project, project_id)
                if cleanup_result['duckdb_tables_deleted'] > 0:
                    logger.info(f"[BACKGROUND] Replaced existing file: cleaned {cleanup_result['duckdb_tables_deleted']} tables")
                
                handler = get_structured_handler()
                
                # ===========================================================
                # PROGRESS CALLBACK - v21 Performance Optimization
                # Maps handler progress (5-95%) to job progress (15-65%)
                # This eliminates the "frozen UI" during large file processing
                # ===========================================================
                def structured_progress_callback(percent: int, message: str):
                    """Map structured handler progress to overall job progress"""
                    # Handler reports 5-95%, we map to 15-65% of overall job
                    # This leaves room for: 0-15% detection, 65-80% intelligence, 80-100% cleanup
                    mapped_percent = 15 + int((percent / 100) * 50)
                    mapped_percent = min(65, max(15, mapped_percent))  # Clamp to 15-65
                    try:
                        ProcessingJobModel.update_progress(job_id, mapped_percent, message)
                    except Exception as cb_e:
                        logger.warning(f"[PROGRESS] Callback error: {cb_e}")
                
                if file_ext == 'csv':
                    result = handler.store_csv(
                        file_path, project, filename,
                        progress_callback=structured_progress_callback,
                        uploaded_by=uploaded_by_email
                    )
                    tables_created = 1
                    total_rows = result.get('row_count', 0)
                else:
                    result = handler.store_excel(
                        file_path, project, filename,
                        progress_callback=structured_progress_callback,
                        uploaded_by=uploaded_by_email
                    )
                    tables_created = len(result.get('tables_created', []))
                    total_rows = result.get('total_rows', 0)
                
                # =====================================================
                # UPLOAD VALIDATION CHECK
                # Warn if data quality issues detected
                # =====================================================
                validation = result.get('validation', {})
                if validation.get('status') in ['warning', 'critical']:
                    issues_count = validation.get('tables_with_issues', 0)
                    issue_tables = [i.get('table', '') for i in validation.get('issues', [])[:3]]
                    warning_msg = f"⚠️ Data quality issues in {issues_count} table(s): {', '.join(issue_tables)}"
                    if issues_count > 3:
                        warning_msg += f" and {issues_count - 3} more"
                    ProcessingJobModel.update_progress(job_id, 68, warning_msg)
                    result['upload_warnings'] = [warning_msg]
                    logger.warning(f"[UPLOAD] Validation issues: {validation}")
                
                ProcessingJobModel.update_progress(
                    job_id, 70, 
                    f"Created {tables_created} table(s) with {total_rows:,} rows"
                )
                
                timing['parse_end'] = time.time()
                timing['storage_start'] = time.time()
                timing['storage_end'] = time.time()  # Storage happens during parse for structured
                
                # =====================================================
                # INTELLIGENCE ANALYSIS - Phase 3
                # Run Tier 1 + 2 analysis on uploaded data
                # =====================================================
                intelligence_summary = run_intelligence_analysis(project, handler, job_id)
                if intelligence_summary:
                    result['intelligence'] = intelligence_summary
                
                # =====================================================
                # DOMAIN INFERENCE - Phase 4
                # Detect what type of data this is for intelligent context
                # =====================================================
                try:
                    from utils.domain_inference_engine import infer_project_domains
                    ProcessingJobModel.update_progress(job_id, 79, "Analyzing data domains...")
                    domains = infer_project_domains(project, project_id, handler)
                    if domains:
                        result['detected_domains'] = domains
                        primary = domains.get('primary_domain')
                        logger.info(f"[BACKGROUND] Domain inference: primary={primary}")
                except ImportError:
                    logger.debug("[BACKGROUND] Domain inference not available")
                except Exception as di_e:
                    logger.warning(f"[BACKGROUND] Domain inference failed: {di_e}")
                
                # Store schema summary in documents table for reference
                if project_id:
                    try:
                        schema_summary = json.dumps(result, indent=2)
                        DocumentModel.create(
                            project_id=project_id,
                            name=filename,
                            category=functional_area or 'Structured Data',
                            file_type=file_ext,
                            file_size=file_size,
                            content=f"STRUCTURED DATA FILE\n\nSchema:\n{schema_summary[:4000]}",
                            metadata={
                                'type': 'structured',
                                'storage': 'duckdb',
                                'tables': result.get('tables_created', []),
                                'total_rows': total_rows,
                                'project': project,
                                'functional_area': functional_area,
                                'intelligence': intelligence_summary
                            }
                        )
                        logger.info(f"[BACKGROUND] Saved structured data metadata to database")
                    except Exception as e:
                        logger.warning(f"[BACKGROUND] Could not save to documents table: {e}")
                
                # Register in document registry with classification (UNIFIED)
                try:
                    times = calc_timing_ms()
                    if REGISTRATION_SERVICE_AVAILABLE:
                        reg_result = RegistrationService.register_structured(
                            filename=filename,
                            project=project,
                            project_id=project_id if not is_global else None,
                            tables_created=result.get('tables_created', []),
                            row_count=total_rows,
                            sheet_count=tables_created,
                            file_content=None,  # Already processed
                            file_size=file_size,
                            file_type=file_ext,
                            uploaded_by_id=uploaded_by_id,
                            uploaded_by_email=uploaded_by_email,
                            job_id=job_id,
                            source=RegistrationSource.UPLOAD,
                            classification_confidence=classification_confidence,
                            content_domain=domain_list,
                            processing_time_ms=times['processing_time_ms'],
                            parse_time_ms=times['parse_time_ms'],
                            storage_time_ms=times['storage_time_ms'],
                            metadata={
                                'project_name': project,
                                'functional_area': functional_area,
                                'file_size': file_size,
                                'intelligence': intelligence_summary,
                                'classification_method': classification_method,
                                'truth_type': truth_type
                            }
                        )
                        if reg_result.success:
                            logger.info(f"[BACKGROUND] Registered: {filename} ({reg_result.lineage_edges} lineage edges)")
                        else:
                            logger.warning(f"[BACKGROUND] Registration warning: {reg_result.error}")
                    else:
                        # Legacy fallback
                        DocumentRegistryModel.register(
                            filename=filename,
                            file_type=file_ext,
                            truth_type=truth_type,
                            classification_method=classification_method,
                            classification_confidence=classification_confidence,
                            content_domain=domain_list,
                            storage_type=DocumentRegistryModel.STORAGE_DUCKDB,
                            project_id=project_id if not is_global else None,
                            is_global=is_global,
                            duckdb_tables=result.get('tables_created', []),
                            row_count=total_rows,
                            sheet_count=tables_created,
                            parse_status='success',
                            processing_time_ms=times['processing_time_ms'],
                            file_hash=file_hash,
                            file_size_bytes=file_size,
                            uploaded_by_id=uploaded_by_id,
                            uploaded_by_email=uploaded_by_email
                        )
                        logger.info(f"[BACKGROUND] Registered (legacy): {filename}")
                    
                except Exception as e:
                    logger.warning(f"[BACKGROUND] Could not register document: {e}")
                
                # Cleanup
                ProcessingJobModel.update_progress(job_id, 90, "Cleaning up...")
                if file_path and os.path.exists(file_path):
                    try:
                        os.remove(file_path)
                    except Exception as e:
                        logger.debug(f"Suppressed: {e}")
                
                # Build completion result
                completion_result = {
                    'filename': filename,
                    'type': 'structured',
                    'tables_created': tables_created,
                    'total_rows': total_rows,
                    'project': project,
                    'functional_area': functional_area
                }
                
                # Add intelligence summary to completion
                if intelligence_summary:
                    completion_result['intelligence'] = intelligence_summary
                
                # Add validation warnings if any
                validation = result.get('validation', {})
                if validation.get('issues'):
                    completion_result['validation'] = validation
                    completion_result['has_data_quality_issues'] = True
                
                # Complete!
                ProcessingJobModel.complete(job_id, completion_result)
                
                # Record metrics for analytics dashboard
                if METRICS_AVAILABLE:
                    times = calc_timing_ms()
                    MetricsService.record_upload(
                        processor='structured',
                        filename=filename,
                        duration_ms=times['processing_time_ms'],
                        success=True,
                        project_id=project_id,
                        file_size_bytes=file_size,
                        rows_processed=total_rows
                    )
                
                logger.info(f"[BACKGROUND] Structured data job {job_id} completed!")
                return
                
            except Exception as e:
                # XLSX/CSV should NEVER go to ChromaDB - structured data doesn't chunk well
                logger.error(f"[BACKGROUND] Structured storage failed: {e}")
                ProcessingJobModel.fail(job_id, f"Failed to process structured data: {str(e)}")
                
                # Record error metric
                if METRICS_AVAILABLE:
                    times = calc_timing_ms()
                    MetricsService.record_upload(
                        processor='structured',
                        filename=filename,
                        duration_ms=times['processing_time_ms'],
                        success=False,
                        project_id=project_id,
                        file_size_bytes=file_size,
                        error_message=str(e)[:500]
                    )
                
                # Cleanup
                if file_path and os.path.exists(file_path):
                    try:
                        os.remove(file_path)
                    except Exception as e:
                        logger.debug(f"Suppressed: {e}")
                return  # Do NOT fall through to RAG
        
        # =================================================================
        # ROUTE 1.5: SMART PDF ROUTING - Check if PDF is tabular
        # =================================================================
        logger.warning(f"[BACKGROUND] Route 1.5 check: file_ext='{file_ext}', SMART_PDF_AVAILABLE={SMART_PDF_AVAILABLE}")
        
        if file_ext == 'pdf' and SMART_PDF_AVAILABLE:
            ProcessingJobModel.update_progress(job_id, 5, "Analyzing PDF structure...")
            logger.warning(f"[BACKGROUND] >>> ENTERING SMART PDF ROUTE for {filename}")
            timing['parse_start'] = time.time()
            
            try:
                def status_callback(msg, progress=None):
                    if progress:
                        # Map smart PDF progress (0-100) to our range (5-50)
                        mapped = 5 + int(progress * 0.45)
                        ProcessingJobModel.update_progress(job_id, mapped, msg)
                    else:
                        logger.warning(f"[SMART-PDF] {msg}")
                
                logger.warning(f"[BACKGROUND] Calling process_pdf_intelligently...")
                
                # Run smart PDF analysis and routing
                pdf_result = process_pdf_intelligently(
                    file_path=file_path,
                    project=project,
                    filename=filename,
                    project_id=project_id,
                    status_callback=status_callback
                )
                
                logger.warning(f"[BACKGROUND] Smart PDF result: success={pdf_result.get('success')}, storage={pdf_result.get('storage_used', [])}")
                
                if pdf_result.get('error'):
                    logger.warning(f"[BACKGROUND] Smart PDF error: {pdf_result.get('error')}")
                
                # Check if DuckDB storage was successful
                duckdb_success = 'duckdb' in pdf_result.get('storage_used', [])
                
                if duckdb_success:
                    # Run intelligence analysis on PDF tables
                    try:
                        handler = get_structured_handler()
                        intelligence_summary = run_intelligence_analysis(project, handler, job_id)
                        if intelligence_summary:
                            pdf_result['intelligence'] = intelligence_summary
                    except Exception as int_e:
                        logger.warning(f"[BACKGROUND] PDF intelligence analysis failed: {int_e}")
                    
                    # DON'T register here - wait until we know if ChromaDB will be used too
                    # Store the duckdb_result info for later registration
                    duckdb_result = pdf_result.get('duckdb_result', {})
                    logger.warning(f"[BACKGROUND] duckdb_result keys: {duckdb_result.keys() if duckdb_result else 'None'}")
                    logger.warning(f"[BACKGROUND] tables_created: {duckdb_result.get('tables_created', [])}")
                    logger.warning(f"[BACKGROUND] total_rows: {duckdb_result.get('total_rows', 0)}")
                
                # Get text content for ChromaDB (either from analysis or re-extract)
                chromadb_result = pdf_result.get('chromadb_result', {})
                text = chromadb_result.get('text_content', '')
                
                if not text:
                    # Fallback: extract text normally
                    text = extract_text(file_path)
                
                # If we got DuckDB storage but no text, we can complete here
                if duckdb_success and (not text or len(text.strip()) < 100):
                    # Register as DuckDB-only (structured)
                    if REGISTRATION_SERVICE_AVAILABLE:
                        try:
                            duckdb_result = pdf_result.get('duckdb_result', {})
                            reg_result = RegistrationService.register_structured(
                                filename=filename,
                                project=project,
                                project_id=project_id,
                                tables_created=duckdb_result.get('tables_created', []),
                                row_count=duckdb_result.get('total_rows', 0),
                                file_size=file_size,
                                file_type='pdf',
                                uploaded_by_id=uploaded_by_id,
                                uploaded_by_email=uploaded_by_email,
                                job_id=job_id,
                                source=RegistrationSource.UPLOAD,
                                classification_confidence=classification_confidence,
                                content_domain=domain_list,
                                metadata={
                                    'project_name': project,
                                    'functional_area': functional_area,
                                    'source': 'smart_pdf',
                                    'exit_reason': 'no_text',
                                    'truth_type': truth_type
                                }
                            )
                            logger.warning(f"[BACKGROUND] Registered PDF (DuckDB only): {reg_result.lineage_edges} lineage edges")
                        except Exception as e:
                            logger.warning(f"[BACKGROUND] Registration failed: {e}")
                    
                    ProcessingJobModel.update_progress(job_id, 90, "Cleaning up...")
                    if file_path and os.path.exists(file_path):
                        try:
                            os.remove(file_path)
                        except Exception as e:
                            logger.debug(f"Suppressed: {e}")
                    
                    ProcessingJobModel.complete(job_id, {
                        'filename': filename,
                        'type': 'smart_pdf',
                        'storage': pdf_result.get('storage_used', []),
                        'duckdb_rows': pdf_result.get('duckdb_result', {}).get('total_rows', 0),
                        'project': project,
                        'intelligence': pdf_result.get('intelligence')
                    })
                    
                    # Record metrics
                    if METRICS_AVAILABLE:
                        times = calc_timing_ms()
                        MetricsService.record_upload(
                            processor='smart_pdf',
                            filename=filename,
                            duration_ms=times['processing_time_ms'],
                            success=True,
                            project_id=project_id,
                            file_size_bytes=file_size,
                            rows_processed=pdf_result.get('duckdb_result', {}).get('total_rows', 0)
                        )
                    
                    logger.info(f"[BACKGROUND] Smart PDF job {job_id} completed (DuckDB only)")
                    return
                
                # Check if smart_pdf_analyzer says to skip ChromaDB (large tabular PDFs)
                skip_chromadb = 'chromadb' not in pdf_result.get('storage_used', [])
                
                if skip_chromadb and duckdb_success:
                    # Register as DuckDB-only (structured) - large tabular
                    if REGISTRATION_SERVICE_AVAILABLE:
                        try:
                            duckdb_result = pdf_result.get('duckdb_result', {})
                            reg_result = RegistrationService.register_structured(
                                filename=filename,
                                project=project,
                                project_id=project_id,
                                tables_created=duckdb_result.get('tables_created', []),
                                row_count=duckdb_result.get('total_rows', 0),
                                file_size=file_size,
                                file_type='pdf',
                                uploaded_by_id=uploaded_by_id,
                                uploaded_by_email=uploaded_by_email,
                                job_id=job_id,
                                source=RegistrationSource.UPLOAD,
                                classification_confidence=classification_confidence,
                                content_domain=domain_list,
                                metadata={
                                    'project_name': project,
                                    'functional_area': functional_area,
                                    'source': 'smart_pdf',
                                    'exit_reason': 'large_tabular_skip_chromadb',
                                    'truth_type': truth_type
                                }
                            )
                            logger.warning(f"[BACKGROUND] Registered PDF (DuckDB only, large tabular): {reg_result.lineage_edges} lineage edges")
                        except Exception as e:
                            logger.warning(f"[BACKGROUND] Registration failed: {e}")
                    
                    # Large tabular PDF - DuckDB only, skip ChromaDB
                    ProcessingJobModel.update_progress(job_id, 90, "Cleaning up (skipping semantic search for large table)...")
                    if file_path and os.path.exists(file_path):
                        try:
                            os.remove(file_path)
                        except Exception as e:
                            logger.debug(f"Suppressed: {e}")
                    
                    ProcessingJobModel.complete(job_id, {
                        'filename': filename,
                        'type': 'smart_pdf_tabular_only',
                        'storage': pdf_result.get('storage_used', []),
                        'duckdb_rows': pdf_result.get('duckdb_result', {}).get('total_rows', 0),
                        'project': project,
                        'note': 'ChromaDB skipped - large tabular PDF',
                        'intelligence': pdf_result.get('intelligence')
                    })
                    logger.info(f"[BACKGROUND] Smart PDF job {job_id} completed (DuckDB only - large tabular)")
                    return
                
                # Continue with ChromaDB using extracted text
                # (Fall through to ROUTE 2 below, but we already have text)
                if text:
                    logger.warning(f"[BACKGROUND] Smart PDF: continuing to ChromaDB with {len(text)} chars")
                    ProcessingJobModel.update_progress(job_id, 50, f"Adding to semantic search ({len(text):,} chars)...")
                
            except Exception as e:
                logger.warning(f"[BACKGROUND] !!! Smart PDF EXCEPTION: {e}")
                import traceback
                logger.warning(traceback.format_exc())
                text = None  # Force re-extraction below
        else:
            text = None
        
        # =================================================================
        # ROUTE 1.75: DOCX/TXT CONTENT ANALYSIS - Route by content, not extension
        # =================================================================
        if file_ext in ['docx', 'txt', 'md'] and DOCUMENT_ANALYZER_AVAILABLE and STRUCTURED_HANDLER_AVAILABLE:
            logger.warning(f"[BACKGROUND] Route 1.75: Analyzing {file_ext} content structure...")
            ProcessingJobModel.update_progress(job_id, 5, "Analyzing document structure...")
            timing['parse_start'] = time.time()
            
            try:
                # Extract text first
                text = extract_text(file_path)
                
                if text and len(text.strip()) > 100:
                    # Analyze content structure
                    analyzer = DocumentAnalyzer()
                    analysis = analyzer.analyze(text, filename, file_ext)
                    
                    logger.warning(f"[BACKGROUND] Content analysis: structure={analysis.structure.value}, "
                                   f"patterns={analysis.patterns}")
                    
                    # If content is TABULAR, route to DuckDB
                    if analysis.structure == DocumentStructure.TABULAR:
                        logger.warning(f"[BACKGROUND] >>> TABULAR content detected in {file_ext}! Routing to DuckDB")
                        ProcessingJobModel.update_progress(job_id, 15, "Detected tabular data - converting for SQL queries...")
                        
                        try:
                            # Convert text to structured format and store in DuckDB
                            handler = get_structured_handler()
                            
                            # Parse the tabular content
                            # Try to detect delimiter and parse as table
                            lines = [l for l in text.strip().split('\n') if l.strip()]
                            
                            if lines:
                                # Detect delimiter
                                first_lines = '\n'.join(lines[:10])
                                if '\t' in first_lines:
                                    delimiter = '\t'
                                elif '|' in first_lines:
                                    delimiter = '|'
                                elif ',' in first_lines and first_lines.count(',') > 3:
                                    delimiter = ','
                                else:
                                    delimiter = None
                                
                                if delimiter:
                                    # Parse as delimited data
                                    import io
                                    
                                    # Clean up pipe-delimited format
                                    if delimiter == '|':
                                        cleaned_lines = []
                                        for line in lines:
                                            # Remove leading/trailing pipes and strip
                                            cleaned = line.strip().strip('|').strip()
                                            if cleaned and not cleaned.startswith('-'):  # Skip separator lines
                                                cleaned_lines.append(cleaned)
                                        text_for_parsing = '\n'.join(cleaned_lines)
                                        delimiter = '|'
                                    else:
                                        text_for_parsing = '\n'.join(lines)
                                    
                                    try:
                                        df = pd.read_csv(io.StringIO(text_for_parsing), sep=delimiter, skipinitialspace=True)
                                        
                                        if len(df) > 0 and len(df.columns) > 1:
                                            # Save as temp CSV and store
                                            temp_csv = f"/tmp/{job_id}_converted.csv"
                                            df.to_csv(temp_csv, index=False)
                                            
                                            result = handler.store_csv(temp_csv, project, f"{filename}_extracted.csv", uploaded_by=uploaded_by_email)
                                            
                                            # Cleanup temp file
                                            try:
                                                os.remove(temp_csv)
                                            except Exception as e:
                                                logger.debug(f"Suppressed: {e}")
                                            
                                            ProcessingJobModel.update_progress(job_id, 70, 
                                                f"Created table with {result.get('row_count', 0):,} rows")
                                            
                                            # Run intelligence analysis
                                            intelligence_summary = run_intelligence_analysis(project, handler, job_id)
                                            
                                            timing['parse_end'] = time.time()
                                            timing['storage_end'] = time.time()
                                            times = calc_timing_ms()
                                            
                                            # Register in document registry with classification (UNIFIED)
                                            try:
                                                if REGISTRATION_SERVICE_AVAILABLE:
                                                    reg_result = RegistrationService.register_structured(
                                                        filename=filename,
                                                        project=project,
                                                        project_id=project_id,
                                                        tables_created=result.get('tables_created', []),
                                                        row_count=result.get('row_count', 0),
                                                        file_size=file_size,
                                                        file_type=file_ext,
                                                        uploaded_by_id=uploaded_by_id,
                                                        uploaded_by_email=uploaded_by_email,
                                                        job_id=job_id,
                                                        source=RegistrationSource.UPLOAD,
                                                        classification_confidence=classification_confidence,
                                                        content_domain=domain_list,
                                                        processing_time_ms=times['processing_time_ms'],
                                                        parse_time_ms=times['parse_time_ms'],
                                                        storage_time_ms=times['storage_time_ms'],
                                                        metadata={
                                                            'project_name': project,
                                                            'functional_area': functional_area,
                                                            'original_type': file_ext,
                                                            'detected_structure': 'tabular',
                                                            'delimiter': delimiter,
                                                            'intelligence': intelligence_summary,
                                                            'truth_type': truth_type
                                                        }
                                                    )
                                                    if not reg_result.success:
                                                        logger.warning(f"[BACKGROUND] Registration warning: {reg_result.error}")
                                                else:
                                                    DocumentRegistryModel.register(
                                                        filename=filename,
                                                        file_type=file_ext,
                                                        truth_type=truth_type,
                                                        classification_method=classification_method,
                                                        classification_confidence=classification_confidence,
                                                        content_domain=domain_list,
                                                        storage_type=DocumentRegistryModel.STORAGE_DUCKDB,
                                                        project_id=project_id,
                                                        is_global=is_global,
                                                        row_count=result.get('row_count', 0),
                                                        parse_status='success',
                                                        processing_time_ms=times['processing_time_ms'],
                                                        file_hash=file_hash,
                                                        file_size_bytes=file_size,
                                                        uploaded_by_id=uploaded_by_id,
                                                        uploaded_by_email=uploaded_by_email
                                                    )
                                            except Exception as reg_e:
                                                logger.warning(f"[BACKGROUND] Could not register: {reg_e}")
                                            
                                            # Cleanup and complete
                                            if file_path and os.path.exists(file_path):
                                                try:
                                                    os.remove(file_path)
                                                except Exception as e:
                                                    logger.debug(f"Suppressed: {e}")
                                            
                                            ProcessingJobModel.complete(job_id, {
                                                'filename': filename,
                                                'type': 'structured_from_text',
                                                'original_format': file_ext,
                                                'rows': result.get('row_count', 0),
                                                'project': project,
                                                'intelligence': intelligence_summary
                                            })
                                            
                                            logger.info(f"[BACKGROUND] Converted {file_ext} to structured data!")
                                            return
                                    
                                    except Exception as parse_e:
                                        logger.warning(f"[BACKGROUND] Could not parse as delimited: {parse_e}")
                                
                                # If delimiter parsing failed, try Word table extraction for DOCX
                                if file_ext == 'docx':
                                    try:
                                        doc = docx.Document(file_path)
                                        tables_data = []
                                        
                                        for table in doc.tables:
                                            table_rows = []
                                            for row in table.rows:
                                                row_data = [cell.text.strip() for cell in row.cells]
                                                table_rows.append(row_data)
                                            
                                            if table_rows:
                                                tables_data.append(table_rows)
                                        
                                        if tables_data:
                                            # Convert first table to DataFrame
                                            first_table = tables_data[0]
                                            if len(first_table) > 1:
                                                df = pd.DataFrame(first_table[1:], columns=first_table[0])
                                                
                                                temp_csv = f"/tmp/{job_id}_docx_table.csv"
                                                df.to_csv(temp_csv, index=False)
                                                
                                                result = handler.store_csv(temp_csv, project, f"{filename}_table.csv", uploaded_by=uploaded_by_email)
                                                
                                                try:
                                                    os.remove(temp_csv)
                                                except Exception as e:
                                                    logger.debug(f"Suppressed: {e}")
                                                
                                                # Run intelligence analysis
                                                intelligence_summary = run_intelligence_analysis(project, handler, job_id)
                                                
                                                # Complete as structured
                                                if file_path and os.path.exists(file_path):
                                                    try:
                                                        os.remove(file_path)
                                                    except Exception as e:
                                                        logger.debug(f"Suppressed: {e}")
                                                
                                                ProcessingJobModel.complete(job_id, {
                                                    'filename': filename,
                                                    'type': 'structured_from_docx_table',
                                                    'tables_found': len(tables_data),
                                                    'rows': result.get('row_count', 0),
                                                    'project': project,
                                                    'intelligence': intelligence_summary
                                                })
                                                
                                                logger.info(f"[BACKGROUND] Extracted {len(tables_data)} tables from DOCX!")
                                                return
                                    
                                    except Exception as docx_e:
                                        logger.warning(f"[BACKGROUND] DOCX table extraction failed: {docx_e}")
                        
                        except Exception as struct_e:
                            logger.warning(f"[BACKGROUND] Structured conversion failed: {struct_e}, continuing to ChromaDB")
                    
                    # If we get here, content is not tabular or conversion failed
                    # Continue to Route 2 with the already-extracted text
                    logger.warning(f"[BACKGROUND] Content is {analysis.structure.value}, continuing to ChromaDB")
                    
            except Exception as analyze_e:
                logger.warning(f"[BACKGROUND] Content analysis failed: {analyze_e}, continuing to ChromaDB")
                text = None  # Force re-extraction
        
        # ROUTE 2: UNSTRUCTURED DATA (PDF/Word/Text) → ChromaDB
        logger.warning(f"[BACKGROUND] Route 2 check: text is {'set' if text else 'None'}")
        
        # Start parse timing if not already started
        if not timing.get('parse_start'):
            timing['parse_start'] = time.time()
        
        if text is None:
            ProcessingJobModel.update_progress(job_id, 5, "Extracting text from file...")
            text = extract_text(file_path)
        
        timing['parse_end'] = time.time()
        
        if not text or len(text.strip()) < 10:
            ProcessingJobModel.fail(job_id, "No text could be extracted from file")
            return
        
        logger.info(f"[BACKGROUND] Extracted {len(text)} characters")
        ProcessingJobModel.update_progress(job_id, 15, f"Extracted {len(text):,} characters")
        
        # Step 2: Prepare metadata (includes truth_type for filtering)
        file_ext = filename.split('.')[-1].lower()
        
        metadata = {
            "project": project,
            "filename": filename,
            "file_type": file_ext,
            "source": filename,
            "upload_date": datetime.now().isoformat(),
            "truth_type": truth_type  # NEW: Enable truth_type filtering in searches
        }
        
        if project_id:
            metadata["project_id"] = project_id
            logger.info(f"[BACKGROUND] ✓ Metadata includes project_id: {project_id}")
        else:
            logger.warning(f"[BACKGROUND] ✗ No project_id in metadata!")
        
        if functional_area:
            metadata["functional_area"] = functional_area
        
        if domain_list:
            metadata["content_domain"] = ','.join(domain_list)
        
        # Step 3: Initialize RAG and process
        ProcessingJobModel.update_progress(job_id, 20, "Initializing document processor...")
        
        # ===========================================================
        # IDEMPOTENT: Clean existing data for this filename first
        # This prevents duplicate chunks if same file uploaded twice
        # BUT: Skip if Smart PDF route already stored to DuckDB!
        # ===========================================================
        if duckdb_success:
            # Smart PDF already stored to DuckDB - only clean ChromaDB chunks, not DuckDB
            logger.info(f"[BACKGROUND] Skipping full cleanup - Smart PDF already stored to DuckDB")
            try:
                rag_cleanup = RAGHandler()
                collection = rag_cleanup.client.get_or_create_collection(name="documents")
                for field in ["filename", "source", "name"]:
                    try:
                        results = collection.get(where={field: filename})
                        if results and results['ids']:
                            collection.delete(ids=results['ids'])
                            logger.info(f"[BACKGROUND] Cleaned {len(results['ids'])} existing ChromaDB chunks")
                            break
                    except Exception as e:
                        logger.debug(f"Suppressed: {e}")
            except Exception as chroma_e:
                logger.warning(f"[BACKGROUND] ChromaDB cleanup error: {chroma_e}")
        else:
            # Normal cleanup - clean both DuckDB and ChromaDB
            cleanup_result = _cleanup_existing_file(filename, project, project_id)
            if cleanup_result['chromadb_chunks_deleted'] > 0:
                logger.info(f"[BACKGROUND] Replaced existing file: cleaned {cleanup_result['chromadb_chunks_deleted']} chunks")
        
        def update_progress(current: int, total: int, message: str):
            """Callback for RAG handler progress updates"""
            # Map RAG progress (0-100) to our range (20-90)
            overall_percent = 20 + int(current * 0.70)
            ProcessingJobModel.update_progress(job_id, overall_percent, message)
        
        rag = RAGHandler()
        
        ProcessingJobModel.update_progress(job_id, 25, "Chunking document...")
        
        timing['embedding_start'] = time.time()
        chunks_added = rag.add_document(
            collection_name="documents",
            text=text,
            metadata=metadata,
            progress_callback=update_progress
        )
        timing['embedding_end'] = time.time()
        timing['storage_end'] = time.time()
        
        if not chunks_added:
            ProcessingJobModel.fail(job_id, "Failed to add document to vector store")
            return
        
        # Step 4: Save to documents table (if we have project UUID)
        ProcessingJobModel.update_progress(job_id, 92, "Saving to database...")
        
        if project_id:
            try:
                DocumentModel.create(
                    project_id=project_id,
                    name=filename,
                    category=functional_area or 'General',
                    file_type=file_ext,
                    file_size=file_size,
                    content=text[:5000],
                    metadata=metadata
                )
                logger.info(f"[BACKGROUND] Saved document to database")
            except Exception as e:
                logger.warning(f"[BACKGROUND] Could not save to documents table: {e}")
        
        # Register in document registry with classification (UNIFIED)
        try:
            times = calc_timing_ms()
            
            if REGISTRATION_SERVICE_AVAILABLE:
                # Use hybrid registration if DuckDB also succeeded
                if duckdb_success:
                    reg_result = RegistrationService.register_hybrid(
                        filename=filename,
                        project=project,
                        project_id=project_id if not is_global else None,
                        tables_created=pdf_result.get('duckdb_result', {}).get('tables_created', []) if 'pdf_result' in dir() else [],
                        row_count=pdf_result.get('duckdb_result', {}).get('total_rows', 0) if 'pdf_result' in dir() else 0,
                        chunk_count=chunks_added,
                        truth_type=truth_type,
                        file_size=file_size,
                        file_type=file_ext,
                        uploaded_by_id=uploaded_by_id,
                        uploaded_by_email=uploaded_by_email,
                        job_id=job_id,
                        source=RegistrationSource.UPLOAD,
                        metadata={
                            'project_name': project,
                            'functional_area': functional_area,
                            'file_size': file_size,
                            'char_count': len(text),
                            'classification_method': classification_method,
                            'classification_confidence': classification_confidence
                        }
                    )
                else:
                    reg_result = RegistrationService.register_embedded(
                        filename=filename,
                        project=project,
                        project_id=project_id if not is_global else None,
                        chunk_count=chunks_added,
                        truth_type=truth_type,
                        file_size=file_size,
                        file_type=file_ext,
                        uploaded_by_id=uploaded_by_id,
                        uploaded_by_email=uploaded_by_email,
                        job_id=job_id,
                        source=RegistrationSource.UPLOAD,
                        classification_confidence=classification_confidence,
                        content_domain=domain_list,
                        functional_area=functional_area,
                        processing_time_ms=times['processing_time_ms'],
                        embedding_time_ms=times['embedding_time_ms'],
                        storage_time_ms=times['storage_time_ms'],
                        metadata={
                            'project_name': project,
                            'functional_area': functional_area,
                            'file_size': file_size,
                            'char_count': len(text),
                            'classification_method': classification_method
                        }
                    )
                
                if reg_result.success:
                    logger.warning(f"[BACKGROUND] Registered: {filename} ({reg_result.lineage_edges} lineage edges)")
                else:
                    logger.warning(f"[BACKGROUND] Registration warning: {reg_result.error}")
            else:
                # Legacy fallback
                storage_type = DocumentRegistryModel.STORAGE_BOTH if duckdb_success else DocumentRegistryModel.STORAGE_CHROMADB
                DocumentRegistryModel.register(
                    filename=filename,
                    file_type=file_ext,
                    truth_type=truth_type,
                    classification_method=classification_method,
                    classification_confidence=classification_confidence,
                    content_domain=domain_list,
                    storage_type=storage_type,
                    project_id=project_id if not is_global else None,
                    is_global=is_global,
                    chunk_count=chunks_added,
                    parse_status='success',
                    processing_time_ms=times['processing_time_ms'],
                    embedding_time_ms=times['embedding_time_ms'],
                    storage_time_ms=times['storage_time_ms'],
                    file_hash=file_hash,
                    file_size_bytes=file_size,
                    uploaded_by_id=uploaded_by_id,
                    uploaded_by_email=uploaded_by_email
                )
                logger.warning(f"[BACKGROUND] Registered (legacy): {filename} as {truth_type}")
            
        except Exception as e:
            logger.warning(f"[BACKGROUND] Could not register document: {e}")
        
        # Cleanup
        ProcessingJobModel.update_progress(job_id, 95, "Cleaning up...")
        if file_path and os.path.exists(file_path):
            try:
                os.remove(file_path)
            except Exception as e:
                logger.debug(f"Suppressed: {e}")
        
        # Complete!
        ProcessingJobModel.complete(job_id, {
            'filename': filename,
            'type': 'unstructured',
            'chunks_created': chunks_added,  # Actual count from RAG
            'project': project,
            'functional_area': functional_area
        })
        
        # Record metrics for analytics dashboard
        if METRICS_AVAILABLE:
            times = calc_timing_ms()
            MetricsService.record_upload(
                processor='semantic',
                filename=filename,
                duration_ms=times['processing_time_ms'],
                success=True,
                project_id=project_id,
                file_size_bytes=file_size,
                chunks_created=chunks_added
            )
        
        logger.info(f"[BACKGROUND] Job {job_id} completed!")
        
    except Exception as e:
        error_msg = f"Processing failed: {str(e)}"
        logger.error(f"[BACKGROUND] {error_msg}")
        logger.error(traceback.format_exc())
        ProcessingJobModel.fail(job_id, error_msg)
        
        # Record error metric
        if METRICS_AVAILABLE:
            try:
                times = calc_timing_ms()
                MetricsService.record_upload(
                    processor='unknown',
                    filename=filename,
                    duration_ms=times['processing_time_ms'],
                    success=False,
                    project_id=project_id,
                    file_size_bytes=file_size,
                    error_message=str(e)[:500]
                )
            except Exception as e:
                logger.debug(f"Suppressed: {e}")  # Don't let metrics fail the error handling
        
        # Cleanup on failure too
        if file_path and os.path.exists(file_path):
            try:
                os.remove(file_path)
            except Exception as e:
                logger.debug(f"Suppressed: {e}")


# =============================================================================
# API ENDPOINTS
# =============================================================================

# =============================================================================
# NOTE: /upload endpoint MOVED to smart_router.py (December 2025)
# =============================================================================
# The unified upload endpoint is now handled by smart_router.py which provides:
# - Single entry point for all file types
# - Automatic routing based on content/truth_type
# - Backward compatibility aliases for /standards/upload, /register/upload
# - Consistent metrics tracking across all upload paths
#
# See: backend/routers/smart_router.py
# =============================================================================

# REMOVED: @router.post("/upload") - now in smart_router.py
# The following function is kept for reference but the endpoint is disabled.
# Smart router calls process_file_background directly for async processing.

# _upload_file_DEPRECATED - Function body removed, logic now in smart_router.py
# See smart_router.smart_upload() for the unified upload endpoint


@router.get("/upload/queue-status")
async def get_queue_status():
    """Get current queue status"""
    return job_queue.get_status()


@router.get("/upload/status/{job_id}")
async def get_job_status(job_id: str):
    """Get processing status for a specific job"""
    try:
        job = ProcessingJobModel.get(job_id)
        if not job:
            raise HTTPException(status_code=404, detail="Job not found")
        
        # Add queue position if still queued
        if job.get('status') in ['pending', 'queued']:
            position = job_queue.get_position(job_id)
            job['queue_position'] = position
        
        return job
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"[STATUS] Error getting job status: {e}")
        raise HTTPException(status_code=500, detail=str(e))


# =============================================================================
# STANDARDS UPLOAD - P4 Standards Layer
# =============================================================================

# Import standards processor
try:
    from backend.utils.standards_processor import (
        process_pdf as standards_process_pdf,
        process_text as standards_process_text,
        get_rule_registry,
        search_standards
    )
    STANDARDS_AVAILABLE = True
except ImportError:
    try:
        from utils.standards_processor import (
            process_pdf as standards_process_pdf,
            process_text as standards_process_text,
            get_rule_registry,
            search_standards
        )
        STANDARDS_AVAILABLE = True
    except ImportError:
        STANDARDS_AVAILABLE = False
        logger.warning("[UPLOAD] Standards processor not available")


@router.get("/standards/health")
async def standards_health():
    """Standards health check."""
    status = {
        "standards_processor": STANDARDS_AVAILABLE,
        "rules_loaded": 0,
        "documents_loaded": 0
    }
    if STANDARDS_AVAILABLE:
        try:
            registry = get_rule_registry()
            status["rules_loaded"] = len(registry.rules)
            status["documents_loaded"] = len(registry.documents)
        except Exception as e:
            logger.debug(f"Suppressed: {e}")
    return status


# =============================================================================
# NOTE: /standards/upload endpoint MOVED to smart_router.py (December 2025)
# =============================================================================
# The unified upload endpoint handles standards via processing_type=standards
# See: backend/routers/smart_router.py
# Backward compatibility alias at /api/standards/upload still works via smart_router
# =============================================================================


@router.get("/standards/rules")
async def list_standards_rules(limit: int = 100):
    """List extracted rules."""
    if not STANDARDS_AVAILABLE:
        raise HTTPException(503, "Standards processor not available")
    registry = get_rule_registry()
    rules = registry.get_all_rules()
    return {"total": len(rules), "rules": [r.to_dict() for r in rules[:limit]]}


@router.get("/standards/documents")
async def list_standards_documents():
    """List processed documents."""
    if not STANDARDS_AVAILABLE:
        raise HTTPException(503, "Standards processor not available")
    registry = get_rule_registry()
    return {"total": len(registry.documents), "documents": [d.to_dict() for d in registry.documents.values()]}


# =============================================================================
# COMPLIANCE CHECK - P4 Standards Layer
# =============================================================================

# Import compliance engine
try:
    from backend.utils.compliance_engine import get_compliance_engine, run_compliance_scan
    COMPLIANCE_ENGINE_AVAILABLE = True
except ImportError:
    try:
        from utils.compliance_engine import get_compliance_engine, run_compliance_scan
        COMPLIANCE_ENGINE_AVAILABLE = True
    except ImportError:
        COMPLIANCE_ENGINE_AVAILABLE = False
        logger.warning("[UPLOAD] Compliance engine not available")


@router.post("/standards/compliance/check/{project_id}")
async def run_compliance_check(
    project_id: str,
    domain: Optional[str] = None
):
    """
    Run compliance check against a project's data.
    
    Uses extracted rules to check for violations in the project's data.
    Returns findings in Five C's format (Condition, Criteria, Cause, Consequence, Corrective Action).
    """
    try:
        logger.info(f"[COMPLIANCE] Starting check for project {project_id}, domain={domain}")
        
        if not STANDARDS_AVAILABLE:
            raise HTTPException(503, "Standards processor not available")
        
        if not COMPLIANCE_ENGINE_AVAILABLE:
            raise HTTPException(503, "Compliance engine not available")
        
        # Get rules
        registry = get_rule_registry()
        
        if domain:
            rules = [r.to_dict() for r in registry.get_rules_by_domain(domain)]
        else:
            rules = [r.to_dict() for r in registry.get_all_rules()]
        
        if not rules:
            return {
                "project_id": project_id,
                "rules_checked": 0,
                "findings": [],
                "findings_count": 0,
                "compliant_count": 0,
                "message": "No rules found. Upload standards documents first."
            }
        
        logger.info(f"[COMPLIANCE] Running {len(rules)} rules against project {project_id}")
        
        # Run compliance scan
        findings = run_compliance_scan(project_id, rules=rules, domain=domain)
        
        logger.info(f"[COMPLIANCE] Scan complete: {len(findings)} findings")
        
        return {
            "project_id": project_id,
            "rules_checked": len(rules),
            "findings_count": len(findings),
            "findings": findings,
            "compliant_count": len(rules) - len(findings)
        }
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"[COMPLIANCE] Check failed: {e}")
        logger.error(traceback.format_exc())
        raise HTTPException(500, f"Compliance check failed: {e}")
