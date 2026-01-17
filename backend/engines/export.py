"""
XLR8 EXPORT ENGINE - Formatted Output Generation
=================================================

Deploy to: backend/engines/export.py

Takes playbook results and generates formatted outputs using
predefined templates. Supports multiple output formats.

EXPORT TEMPLATES:
- executive_summary: High-level overview for leadership
- findings_report: Detailed findings with evidence
- data_quality_scorecard: Pass/fail metrics grid
- remediation_checklist: Action items from findings
- raw_data: CSV/JSON dump of all results

OUTPUT FORMATS:
- PDF (via reportlab or weasyprint)
- DOCX (via python-docx)
- XLSX (via openpyxl)
- CSV (native)
- JSON (native)
- HTML (native - can convert to PDF)

Created: January 17, 2026
"""

import logging
import json
import csv
import io
from typing import Dict, List, Optional, Any
from dataclasses import dataclass, field
from datetime import datetime
from enum import Enum

logger = logging.getLogger(__name__)


# =============================================================================
# ENUMS AND DATA CLASSES
# =============================================================================

class ExportFormat(Enum):
    """Supported export formats."""
    PDF = "pdf"
    DOCX = "docx"
    XLSX = "xlsx"
    CSV = "csv"
    JSON = "json"
    HTML = "html"


class TemplateType(Enum):
    """Pre-defined export templates."""
    EXECUTIVE_SUMMARY = "executive_summary"
    FINDINGS_REPORT = "findings_report"
    DATA_QUALITY_SCORECARD = "data_quality_scorecard"
    REMEDIATION_CHECKLIST = "remediation_checklist"
    RAW_DATA = "raw_data"


@dataclass
class ExportTemplate:
    """Definition of an export template."""
    id: str
    name: str
    description: str
    template_type: TemplateType
    default_format: ExportFormat
    supported_formats: List[ExportFormat]
    sections: List[Dict[str, Any]]  # Template sections with placeholders
    

@dataclass
class ExportResult:
    """Result from export operation."""
    success: bool
    template: str
    format: str
    filename: str
    content: bytes  # Raw file content
    content_type: str  # MIME type
    metadata: Dict[str, Any] = field(default_factory=dict)
    error: Optional[str] = None


# =============================================================================
# TEMPLATE LIBRARY
# =============================================================================

EXPORT_TEMPLATES = {
    TemplateType.EXECUTIVE_SUMMARY: ExportTemplate(
        id="executive_summary",
        name="Executive Summary",
        description="High-level overview for leadership - key metrics and critical findings only",
        template_type=TemplateType.EXECUTIVE_SUMMARY,
        default_format=ExportFormat.PDF,
        supported_formats=[ExportFormat.PDF, ExportFormat.DOCX, ExportFormat.HTML],
        sections=[
            {
                "id": "header",
                "title": "Analysis Summary",
                "type": "header",
                "fields": ["project_name", "run_date", "playbook_name"]
            },
            {
                "id": "metrics",
                "title": "Key Metrics",
                "type": "metrics_grid",
                "fields": [
                    {"label": "Features Run", "value": "{{total_features}}"},
                    {"label": "Passed", "value": "{{passed_count}}", "color": "green"},
                    {"label": "Findings", "value": "{{total_findings}}", "color": "red"},
                    {"label": "Data Quality Score", "value": "{{quality_score}}%"}
                ]
            },
            {
                "id": "critical_findings",
                "title": "Critical Findings",
                "type": "findings_list",
                "filter": {"severity": ["critical", "error"]},
                "max_items": 5
            },
            {
                "id": "recommendations",
                "title": "Recommended Actions",
                "type": "recommendations",
                "max_items": 3
            }
        ]
    ),
    
    TemplateType.FINDINGS_REPORT: ExportTemplate(
        id="findings_report",
        name="Detailed Findings Report",
        description="Comprehensive report with all findings, evidence, and remediation guidance",
        template_type=TemplateType.FINDINGS_REPORT,
        default_format=ExportFormat.DOCX,
        supported_formats=[ExportFormat.DOCX, ExportFormat.PDF, ExportFormat.HTML],
        sections=[
            {
                "id": "cover",
                "title": "{{playbook_name}} - Analysis Report",
                "type": "cover_page",
                "fields": ["project_name", "client_name", "run_date", "prepared_by"]
            },
            {
                "id": "toc",
                "title": "Table of Contents",
                "type": "table_of_contents"
            },
            {
                "id": "executive_summary",
                "title": "Executive Summary",
                "type": "text_section",
                "content": "{{executive_summary}}"
            },
            {
                "id": "methodology",
                "title": "Methodology",
                "type": "text_section",
                "content": "This analysis was performed using XLR8's {{engine_count}} universal analysis engines..."
            },
            {
                "id": "findings_by_category",
                "title": "Findings by Category",
                "type": "findings_grouped",
                "group_by": "category",
                "include_evidence": True
            },
            {
                "id": "data_summary",
                "title": "Data Summary",
                "type": "data_table",
                "fields": ["table_name", "row_count", "findings_count"]
            },
            {
                "id": "appendix",
                "title": "Appendix: Raw Results",
                "type": "raw_results",
                "include_sql": True
            }
        ]
    ),
    
    TemplateType.DATA_QUALITY_SCORECARD: ExportTemplate(
        id="data_quality_scorecard",
        name="Data Quality Scorecard",
        description="Pass/fail grid showing quality metrics by category and table",
        template_type=TemplateType.DATA_QUALITY_SCORECARD,
        default_format=ExportFormat.XLSX,
        supported_formats=[ExportFormat.XLSX, ExportFormat.PDF, ExportFormat.HTML, ExportFormat.CSV],
        sections=[
            {
                "id": "summary_sheet",
                "title": "Summary",
                "type": "summary_metrics",
                "fields": ["overall_score", "by_category", "by_engine", "trend"]
            },
            {
                "id": "detail_sheet",
                "title": "Detail by Check",
                "type": "checks_grid",
                "columns": ["check_name", "engine", "table", "status", "findings", "score"]
            },
            {
                "id": "by_table_sheet",
                "title": "By Table",
                "type": "table_scores",
                "columns": ["table_name", "checks_run", "passed", "failed", "score"]
            }
        ]
    ),
    
    TemplateType.REMEDIATION_CHECKLIST: ExportTemplate(
        id="remediation_checklist",
        name="Remediation Checklist",
        description="Actionable checklist of items to fix, grouped by priority",
        template_type=TemplateType.REMEDIATION_CHECKLIST,
        default_format=ExportFormat.XLSX,
        supported_formats=[ExportFormat.XLSX, ExportFormat.DOCX, ExportFormat.CSV],
        sections=[
            {
                "id": "critical",
                "title": "Critical - Fix Immediately",
                "type": "checklist",
                "filter": {"severity": "critical"},
                "columns": ["checkbox", "finding", "table", "affected_rows", "remediation"]
            },
            {
                "id": "high",
                "title": "High Priority",
                "type": "checklist",
                "filter": {"severity": "error"},
                "columns": ["checkbox", "finding", "table", "affected_rows", "remediation"]
            },
            {
                "id": "medium",
                "title": "Medium Priority",
                "type": "checklist",
                "filter": {"severity": "warning"},
                "columns": ["checkbox", "finding", "table", "affected_rows", "remediation"]
            },
            {
                "id": "low",
                "title": "Low Priority / Informational",
                "type": "checklist",
                "filter": {"severity": "info"},
                "columns": ["checkbox", "finding", "table", "affected_rows", "remediation"]
            }
        ]
    ),
    
    TemplateType.RAW_DATA: ExportTemplate(
        id="raw_data",
        name="Raw Data Export",
        description="Complete data dump of all results for further analysis",
        template_type=TemplateType.RAW_DATA,
        default_format=ExportFormat.JSON,
        supported_formats=[ExportFormat.JSON, ExportFormat.CSV, ExportFormat.XLSX],
        sections=[
            {
                "id": "results",
                "title": "All Results",
                "type": "raw_dump",
                "include": ["metadata", "findings", "data", "sql"]
            }
        ]
    )
}


# =============================================================================
# EXPORT ENGINE
# =============================================================================

class ExportEngine:
    """
    Generates formatted exports from playbook results.
    
    Usage:
        engine = ExportEngine()
        result = engine.export(
            playbook_results=batch_results,
            template=TemplateType.EXECUTIVE_SUMMARY,
            format=ExportFormat.PDF,
            context={"project_name": "Acme Corp", "client_name": "Acme"}
        )
        
        # result.content contains the file bytes
        # result.filename is the suggested filename
        # result.content_type is the MIME type
    """
    
    def __init__(self):
        self.templates = EXPORT_TEMPLATES
    
    def get_available_templates(self) -> List[Dict]:
        """Get list of available export templates."""
        return [
            {
                "id": t.id,
                "name": t.name,
                "description": t.description,
                "default_format": t.default_format.value,
                "supported_formats": [f.value for f in t.supported_formats]
            }
            for t in self.templates.values()
        ]
    
    def export(
        self,
        playbook_results: Dict,
        template: TemplateType,
        format: ExportFormat = None,
        context: Dict = None
    ) -> ExportResult:
        """
        Generate export from playbook results.
        
        Args:
            playbook_results: Results from /api/engines/{project}/batch
            template: Which template to use
            format: Output format (defaults to template's default)
            context: Additional context (project_name, client_name, etc.)
            
        Returns:
            ExportResult with file content
        """
        try:
            template_def = self.templates.get(template)
            if not template_def:
                return ExportResult(
                    success=False,
                    template=template.value if template else "unknown",
                    format="",
                    filename="",
                    content=b"",
                    content_type="",
                    error=f"Unknown template: {template}"
                )
            
            # Use default format if not specified
            if format is None:
                format = template_def.default_format
            
            # Check format is supported
            if format not in template_def.supported_formats:
                return ExportResult(
                    success=False,
                    template=template_def.id,
                    format=format.value,
                    filename="",
                    content=b"",
                    content_type="",
                    error=f"Format {format.value} not supported for template {template_def.id}"
                )
            
            # Build export context
            export_context = self._build_context(playbook_results, context or {})
            
            # Generate based on format
            if format == ExportFormat.JSON:
                return self._export_json(template_def, export_context)
            elif format == ExportFormat.CSV:
                return self._export_csv(template_def, export_context)
            elif format == ExportFormat.HTML:
                return self._export_html(template_def, export_context)
            elif format == ExportFormat.XLSX:
                return self._export_xlsx(template_def, export_context)
            elif format == ExportFormat.DOCX:
                return self._export_docx(template_def, export_context)
            elif format == ExportFormat.PDF:
                return self._export_pdf(template_def, export_context)
            else:
                return ExportResult(
                    success=False,
                    template=template_def.id,
                    format=format.value,
                    filename="",
                    content=b"",
                    content_type="",
                    error=f"Format {format.value} not yet implemented"
                )
                
        except Exception as e:
            logger.error(f"[EXPORT] Error: {e}")
            import traceback
            return ExportResult(
                success=False,
                template=template.value if template else "unknown",
                format=format.value if format else "unknown",
                filename="",
                content=b"",
                content_type="",
                error=str(e),
                metadata={"traceback": traceback.format_exc()}
            )
    
    def _build_context(self, playbook_results: Dict, user_context: Dict) -> Dict:
        """Build the full context for template rendering."""
        # Extract metrics from results
        results = playbook_results.get("results", [])
        
        total_features = len(results)
        passed_count = sum(1 for r in results if r.get("success"))
        failed_count = total_features - passed_count
        
        all_findings = []
        for r in results:
            findings = r.get("findings", [])
            for f in findings:
                f["source_engine"] = r.get("engine")
                f["source_feature"] = r.get("id")
            all_findings.extend(findings)
        
        total_findings = len(all_findings)
        
        # Calculate quality score (simple: % passed)
        quality_score = round((passed_count / total_features * 100) if total_features > 0 else 0)
        
        # Group findings by severity
        findings_by_severity = {}
        for f in all_findings:
            sev = f.get("severity", "info")
            if sev not in findings_by_severity:
                findings_by_severity[sev] = []
            findings_by_severity[sev].append(f)
        
        # Group findings by category/type
        findings_by_type = {}
        for f in all_findings:
            ftype = f.get("finding_type", "unknown")
            if ftype not in findings_by_type:
                findings_by_type[ftype] = []
            findings_by_type[ftype].append(f)
        
        return {
            # User-provided context
            **user_context,
            
            # Computed metrics
            "run_date": datetime.now().strftime("%Y-%m-%d %H:%M"),
            "total_features": total_features,
            "passed_count": passed_count,
            "failed_count": failed_count,
            "total_findings": total_findings,
            "quality_score": quality_score,
            
            # Grouped data
            "findings_by_severity": findings_by_severity,
            "findings_by_type": findings_by_type,
            "all_findings": all_findings,
            
            # Raw results
            "results": results,
            "raw": playbook_results
        }
    
    def _export_json(self, template: ExportTemplate, context: Dict) -> ExportResult:
        """Export as JSON."""
        filename = f"{template.id}_{context.get('project_name', 'export')}_{datetime.now().strftime('%Y%m%d')}.json"
        
        # For raw data, include everything
        if template.template_type == TemplateType.RAW_DATA:
            content = json.dumps(context["raw"], indent=2, default=str)
        else:
            # For other templates, include context minus raw
            export_data = {k: v for k, v in context.items() if k != "raw"}
            export_data["template"] = template.id
            content = json.dumps(export_data, indent=2, default=str)
        
        return ExportResult(
            success=True,
            template=template.id,
            format="json",
            filename=filename,
            content=content.encode("utf-8"),
            content_type="application/json",
            metadata={"record_count": len(context.get("all_findings", []))}
        )
    
    def _export_csv(self, template: ExportTemplate, context: Dict) -> ExportResult:
        """Export as CSV."""
        filename = f"{template.id}_{context.get('project_name', 'export')}_{datetime.now().strftime('%Y%m%d')}.csv"
        
        output = io.StringIO()
        
        # For findings-based templates, export findings
        findings = context.get("all_findings", [])
        
        if findings:
            # Get all unique keys from findings
            all_keys = set()
            for f in findings:
                all_keys.update(f.keys())
            
            fieldnames = sorted(all_keys)
            writer = csv.DictWriter(output, fieldnames=fieldnames, extrasaction='ignore')
            writer.writeheader()
            writer.writerows(findings)
        else:
            # Export results summary
            results = context.get("results", [])
            if results:
                fieldnames = ["id", "engine", "success", "status", "row_count", "summary", "findings_count"]
                writer = csv.DictWriter(output, fieldnames=fieldnames)
                writer.writeheader()
                for r in results:
                    writer.writerow({
                        "id": r.get("id"),
                        "engine": r.get("engine"),
                        "success": r.get("success"),
                        "status": r.get("status"),
                        "row_count": r.get("row_count"),
                        "summary": r.get("summary"),
                        "findings_count": len(r.get("findings", []))
                    })
        
        content = output.getvalue()
        
        return ExportResult(
            success=True,
            template=template.id,
            format="csv",
            filename=filename,
            content=content.encode("utf-8"),
            content_type="text/csv",
            metadata={"record_count": len(findings) or len(context.get("results", []))}
        )
    
    def _export_html(self, template: ExportTemplate, context: Dict) -> ExportResult:
        """Export as HTML."""
        filename = f"{template.id}_{context.get('project_name', 'export')}_{datetime.now().strftime('%Y%m%d')}.html"
        
        html = self._render_html_template(template, context)
        
        return ExportResult(
            success=True,
            template=template.id,
            format="html",
            filename=filename,
            content=html.encode("utf-8"),
            content_type="text/html",
            metadata={}
        )
    
    def _render_html_template(self, template: ExportTemplate, context: Dict) -> str:
        """Render HTML from template."""
        # Build HTML based on template sections
        sections_html = []
        
        for section in template.sections:
            section_type = section.get("type")
            
            if section_type == "header":
                sections_html.append(f"""
                    <div class="header">
                        <h1>{template.name}</h1>
                        <p>Project: {context.get('project_name', 'N/A')} | Date: {context.get('run_date')}</p>
                    </div>
                """)
            
            elif section_type == "metrics_grid":
                metrics_html = '<div class="metrics-grid">'
                for field in section.get("fields", []):
                    label = field.get("label", "")
                    value_template = field.get("value", "")
                    # Simple template replacement
                    value = value_template.replace("{{total_features}}", str(context.get("total_features", 0)))
                    value = value.replace("{{passed_count}}", str(context.get("passed_count", 0)))
                    value = value.replace("{{total_findings}}", str(context.get("total_findings", 0)))
                    value = value.replace("{{quality_score}}", str(context.get("quality_score", 0)))
                    
                    color = field.get("color", "")
                    color_class = f"metric-{color}" if color else ""
                    metrics_html += f'<div class="metric {color_class}"><div class="metric-value">{value}</div><div class="metric-label">{label}</div></div>'
                metrics_html += '</div>'
                sections_html.append(metrics_html)
            
            elif section_type == "findings_list":
                findings = context.get("all_findings", [])
                filter_severities = section.get("filter", {}).get("severity", [])
                if filter_severities:
                    findings = [f for f in findings if f.get("severity") in filter_severities]
                
                max_items = section.get("max_items", 10)
                findings = findings[:max_items]
                
                findings_html = f'<div class="section"><h2>{section.get("title", "Findings")}</h2>'
                if findings:
                    findings_html += '<ul class="findings-list">'
                    for f in findings:
                        findings_html += f'<li class="finding finding-{f.get("severity", "info")}"><strong>{f.get("finding_type")}</strong>: {f.get("message")}</li>'
                    findings_html += '</ul>'
                else:
                    findings_html += '<p class="no-findings">No findings in this category.</p>'
                findings_html += '</div>'
                sections_html.append(findings_html)
            
            elif section_type == "findings_grouped":
                findings_by_type = context.get("findings_by_type", {})
                grouped_html = f'<div class="section"><h2>{section.get("title", "Findings")}</h2>'
                
                for ftype, findings in findings_by_type.items():
                    grouped_html += f'<h3>{ftype} ({len(findings)})</h3>'
                    grouped_html += '<ul class="findings-list">'
                    for f in findings[:10]:  # Limit per group
                        grouped_html += f'<li class="finding finding-{f.get("severity", "info")}">{f.get("message")}</li>'
                    if len(findings) > 10:
                        grouped_html += f'<li class="more">... and {len(findings) - 10} more</li>'
                    grouped_html += '</ul>'
                
                grouped_html += '</div>'
                sections_html.append(grouped_html)
        
        # Wrap in full HTML document
        html = f"""<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <title>{template.name} - {context.get('project_name', 'Export')}</title>
    <style>
        body {{ font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, sans-serif; margin: 40px; color: #1a1a2e; }}
        .header {{ border-bottom: 2px solid #83b16d; padding-bottom: 20px; margin-bottom: 30px; }}
        .header h1 {{ margin: 0 0 10px; color: #1a1a2e; }}
        .header p {{ margin: 0; color: #64748b; }}
        .metrics-grid {{ display: grid; grid-template-columns: repeat(4, 1fr); gap: 20px; margin: 30px 0; }}
        .metric {{ background: #f8f9fa; padding: 20px; border-radius: 8px; text-align: center; }}
        .metric-value {{ font-size: 32px; font-weight: 600; }}
        .metric-label {{ font-size: 14px; color: #64748b; margin-top: 5px; }}
        .metric-green .metric-value {{ color: #16a34a; }}
        .metric-red .metric-value {{ color: #dc2626; }}
        .section {{ margin: 30px 0; }}
        .section h2 {{ color: #1a1a2e; border-bottom: 1px solid #e2e8f0; padding-bottom: 10px; }}
        .section h3 {{ color: #64748b; font-size: 16px; }}
        .findings-list {{ list-style: none; padding: 0; }}
        .finding {{ padding: 12px; margin: 8px 0; border-radius: 6px; border-left: 4px solid; }}
        .finding-critical {{ background: #fef2f2; border-color: #dc2626; }}
        .finding-error {{ background: #fef2f2; border-color: #dc2626; }}
        .finding-warning {{ background: #fffbeb; border-color: #d97706; }}
        .finding-info {{ background: #f0f9ff; border-color: #0284c7; }}
        .no-findings {{ color: #64748b; font-style: italic; }}
        .more {{ color: #64748b; font-style: italic; }}
    </style>
</head>
<body>
    {''.join(sections_html)}
</body>
</html>"""
        
        return html
    
    def _export_xlsx(self, template: ExportTemplate, context: Dict) -> ExportResult:
        """Export as XLSX."""
        try:
            import openpyxl
            from openpyxl.styles import Font, PatternFill, Alignment
        except ImportError:
            return ExportResult(
                success=False,
                template=template.id,
                format="xlsx",
                filename="",
                content=b"",
                content_type="",
                error="openpyxl not installed - cannot generate XLSX"
            )
        
        filename = f"{template.id}_{context.get('project_name', 'export')}_{datetime.now().strftime('%Y%m%d')}.xlsx"
        
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "Summary"
        
        # Header styling
        header_font = Font(bold=True, color="FFFFFF")
        header_fill = PatternFill(start_color="83b16d", end_color="83b16d", fill_type="solid")
        
        # Summary sheet
        ws["A1"] = "Metric"
        ws["B1"] = "Value"
        ws["A1"].font = header_font
        ws["A1"].fill = header_fill
        ws["B1"].font = header_font
        ws["B1"].fill = header_fill
        
        metrics = [
            ("Project", context.get("project_name", "N/A")),
            ("Run Date", context.get("run_date", "N/A")),
            ("Total Features", context.get("total_features", 0)),
            ("Passed", context.get("passed_count", 0)),
            ("Failed", context.get("failed_count", 0)),
            ("Total Findings", context.get("total_findings", 0)),
            ("Quality Score", f"{context.get('quality_score', 0)}%")
        ]
        
        for i, (label, value) in enumerate(metrics, start=2):
            ws[f"A{i}"] = label
            ws[f"B{i}"] = value
        
        # Findings sheet
        ws_findings = wb.create_sheet("Findings")
        findings = context.get("all_findings", [])
        
        if findings:
            headers = ["Severity", "Type", "Message", "Engine", "Feature"]
            for col, header in enumerate(headers, start=1):
                cell = ws_findings.cell(row=1, column=col, value=header)
                cell.font = header_font
                cell.fill = header_fill
            
            for row, f in enumerate(findings, start=2):
                ws_findings.cell(row=row, column=1, value=f.get("severity", ""))
                ws_findings.cell(row=row, column=2, value=f.get("finding_type", ""))
                ws_findings.cell(row=row, column=3, value=f.get("message", ""))
                ws_findings.cell(row=row, column=4, value=f.get("source_engine", ""))
                ws_findings.cell(row=row, column=5, value=f.get("source_feature", ""))
        
        # Results sheet
        ws_results = wb.create_sheet("Results")
        results = context.get("results", [])
        
        if results:
            headers = ["Feature ID", "Engine", "Success", "Status", "Rows", "Findings", "Summary"]
            for col, header in enumerate(headers, start=1):
                cell = ws_results.cell(row=1, column=col, value=header)
                cell.font = header_font
                cell.fill = header_fill
            
            for row, r in enumerate(results, start=2):
                ws_results.cell(row=row, column=1, value=r.get("id", ""))
                ws_results.cell(row=row, column=2, value=r.get("engine", ""))
                ws_results.cell(row=row, column=3, value="Yes" if r.get("success") else "No")
                ws_results.cell(row=row, column=4, value=r.get("status", ""))
                ws_results.cell(row=row, column=5, value=r.get("row_count", 0))
                ws_results.cell(row=row, column=6, value=len(r.get("findings", [])))
                ws_results.cell(row=row, column=7, value=r.get("summary", ""))
        
        # Save to bytes
        output = io.BytesIO()
        wb.save(output)
        content = output.getvalue()
        
        return ExportResult(
            success=True,
            template=template.id,
            format="xlsx",
            filename=filename,
            content=content,
            content_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            metadata={"sheets": ["Summary", "Findings", "Results"]}
        )
    
    def _export_docx(self, template: ExportTemplate, context: Dict) -> ExportResult:
        """Export as DOCX."""
        try:
            from docx import Document
            from docx.shared import Inches, Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError:
            return ExportResult(
                success=False,
                template=template.id,
                format="docx",
                filename="",
                content=b"",
                content_type="",
                error="python-docx not installed - cannot generate DOCX"
            )
        
        filename = f"{template.id}_{context.get('project_name', 'export')}_{datetime.now().strftime('%Y%m%d')}.docx"
        
        doc = Document()
        
        # Title
        title = doc.add_heading(template.name, 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Subtitle
        subtitle = doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = subtitle.add_run(f"Project: {context.get('project_name', 'N/A')} | Date: {context.get('run_date')}")
        run.font.size = Pt(12)
        
        doc.add_paragraph()
        
        # Summary section
        doc.add_heading("Summary", level=1)
        
        summary_table = doc.add_table(rows=4, cols=2)
        summary_table.style = 'Table Grid'
        
        metrics = [
            ("Features Run", str(context.get("total_features", 0))),
            ("Passed", str(context.get("passed_count", 0))),
            ("Total Findings", str(context.get("total_findings", 0))),
            ("Quality Score", f"{context.get('quality_score', 0)}%")
        ]
        
        for i, (label, value) in enumerate(metrics):
            summary_table.rows[i].cells[0].text = label
            summary_table.rows[i].cells[1].text = value
        
        doc.add_paragraph()
        
        # Findings section
        findings = context.get("all_findings", [])
        if findings:
            doc.add_heading("Findings", level=1)
            
            for f in findings[:20]:  # Limit to 20
                p = doc.add_paragraph()
                severity = f.get("severity", "info").upper()
                p.add_run(f"[{severity}] ").bold = True
                p.add_run(f"{f.get('finding_type', 'Unknown')}: ")
                p.add_run(f.get("message", ""))
        
        # Save to bytes
        output = io.BytesIO()
        doc.save(output)
        content = output.getvalue()
        
        return ExportResult(
            success=True,
            template=template.id,
            format="docx",
            filename=filename,
            content=content,
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            metadata={"findings_count": len(findings)}
        )
    
    def _export_pdf(self, template: ExportTemplate, context: Dict) -> ExportResult:
        """Export as PDF (via HTML)."""
        # For now, generate HTML and note that PDF conversion would need weasyprint/puppeteer
        html_result = self._export_html(template, context)
        
        if not html_result.success:
            return html_result
        
        # Try to convert HTML to PDF if weasyprint available
        try:
            from weasyprint import HTML
            
            filename = f"{template.id}_{context.get('project_name', 'export')}_{datetime.now().strftime('%Y%m%d')}.pdf"
            
            pdf_content = HTML(string=html_result.content.decode('utf-8')).write_pdf()
            
            return ExportResult(
                success=True,
                template=template.id,
                format="pdf",
                filename=filename,
                content=pdf_content,
                content_type="application/pdf",
                metadata={}
            )
        except ImportError:
            # Return HTML with note about PDF
            return ExportResult(
                success=True,
                template=template.id,
                format="html",  # Fallback to HTML
                filename=html_result.filename.replace('.html', '_for_pdf.html'),
                content=html_result.content,
                content_type="text/html",
                metadata={"note": "PDF generation requires weasyprint - returning HTML instead"}
            )


# =============================================================================
# CONVENIENCE FUNCTIONS
# =============================================================================

def get_export_templates() -> List[Dict]:
    """Get list of available export templates."""
    engine = ExportEngine()
    return engine.get_available_templates()


def export_playbook_results(
    playbook_results: Dict,
    template: str,
    format: str = None,
    context: Dict = None
) -> ExportResult:
    """
    Convenience function to export playbook results.
    
    Args:
        playbook_results: Results from /api/engines/{project}/batch
        template: Template name (executive_summary, findings_report, etc.)
        format: Output format (pdf, docx, xlsx, csv, json, html)
        context: Additional context (project_name, client_name, etc.)
    """
    engine = ExportEngine()
    
    # Convert string to enum
    try:
        template_enum = TemplateType(template)
    except ValueError:
        return ExportResult(
            success=False,
            template=template,
            format=format or "",
            filename="",
            content=b"",
            content_type="",
            error=f"Unknown template: {template}. Available: {[t.value for t in TemplateType]}"
        )
    
    format_enum = None
    if format:
        try:
            format_enum = ExportFormat(format)
        except ValueError:
            return ExportResult(
                success=False,
                template=template,
                format=format,
                filename="",
                content=b"",
                content_type="",
                error=f"Unknown format: {format}. Available: {[f.value for f in ExportFormat]}"
            )
    
    return engine.export(playbook_results, template_enum, format_enum, context)
