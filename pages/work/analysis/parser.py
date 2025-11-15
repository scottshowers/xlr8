"""
Document Parser Module
Uses the ACTUAL working EnhancedPayrollParser from secure_pdf_parser.py
"""

import streamlit as st
from typing import Dict, Any, Optional
from utils.secure_pdf_parser import EnhancedPayrollParser


def parse_document(uploaded_file) -> Optional[Dict[str, Any]]:
    """
    Parse uploaded document using ACTUAL working parser
    
    Args:
        uploaded_file: Streamlit UploadedFile object
    
    Returns:
        Dictionary with parsed data including tables and metadata
    """
    
    try:
        file_type = uploaded_file.name.split('.')[-1].lower()
        
        if file_type == 'pdf':
            return _parse_pdf(uploaded_file)
        elif file_type in ['xlsx', 'xls', 'csv']:
            return _parse_excel(uploaded_file)
        elif file_type in ['docx', 'doc']:
            return _parse_word(uploaded_file)
        else:
            st.error(f"Unsupported file type: {file_type}")
            return None
            
    except Exception as e:
        st.error(f"Error parsing document: {str(e)}")
        return None


def _parse_pdf(uploaded_file) -> Dict[str, Any]:
    """Parse PDF using ACTUAL working EnhancedPayrollParser"""
    
    # Get parser from session state or create new one
    if 'pdf_parser' not in st.session_state:
        st.session_state.pdf_parser = EnhancedPayrollParser()
    
    parser = st.session_state.pdf_parser
    
    # Parse the PDF using the working parser
    result = parser.parse_pdf(
        pdf_content=uploaded_file.read(),
        filename=uploaded_file.name,
        pages='all'
    )
    
    if result.get('success') and result.get('tables'):
        # Convert to format expected by analysis module
        text_content = ""
        for table_info in result['tables']:
            df = table_info['data']
            text_content += f"\n\nTable {table_info['table_num']}:\n"
            text_content += df.to_string()
        
        return {
            'text': text_content,
            'tables': result['tables'],
            'metadata': {
                'type': 'pdf',
                'filename': uploaded_file.name,
                'method': result.get('method', 'unknown'),
                'processing_time': result.get('processing_time', 0),
                'total_pages': result.get('metadata', {}).get('total_pages', 0)
            },
            'raw_result': result  # Keep full result for advanced processing
        }
    else:
        # Parser failed
        error_msg = result.get('error', 'Failed to parse PDF')
        st.error(f"PDF parsing failed: {error_msg}")
        return None


def _parse_excel(uploaded_file) -> Dict[str, Any]:
    """Parse Excel document"""
    import pandas as pd
    
    try:
        # Read all sheets
        excel_file = pd.ExcelFile(uploaded_file)
        sheets_data = {}
        text_content = ""
        
        for sheet_name in excel_file.sheet_names:
            df = pd.read_excel(excel_file, sheet_name=sheet_name)
            sheets_data[sheet_name] = df
            text_content += f"\n\nSheet: {sheet_name}\n{df.to_string()}\n"
        
        return {
            'text': text_content,
            'tables': [{'name': name, 'data': df} for name, df in sheets_data.items()],
            'metadata': {
                'type': 'excel',
                'filename': uploaded_file.name,
                'sheets': list(sheets_data.keys()),
                'total_sheets': len(sheets_data)
            }
        }
    except Exception as e:
        st.error(f"Error parsing Excel: {str(e)}")
        return None


def _parse_word(uploaded_file) -> Dict[str, Any]:
    """Parse Word document"""
    try:
        from docx import Document
        
        doc = Document(uploaded_file)
        text_content = "\n".join([para.text for para in doc.paragraphs])
        
        return {
            'text': text_content,
            'tables': [],
            'metadata': {
                'type': 'word',
                'filename': uploaded_file.name,
                'paragraphs': len(doc.paragraphs)
            }
        }
    except Exception as e:
        st.error(f"Error parsing Word: {str(e)}")
        return None


# Standalone testing
if __name__ == "__main__":
    st.title("Parser Module - Using ACTUAL Working Parser")
    st.write("Upload a test file to verify parsing with EnhancedPayrollParser")
    
    test_file = st.file_uploader("Test file", type=['pdf', 'xlsx', 'csv'])
    
    if test_file:
        with st.spinner("Parsing..."):
            result = parse_document(test_file)
        
        if result:
            st.success("✅ Parse successful!")
            st.json(result['metadata'])
            
            if result.get('tables'):
                st.write(f"Found {len(result['tables'])} tables")
                for i, table in enumerate(result['tables'][:3]):  # Show first 3
                    st.write(f"Table {i+1}:")
                    st.dataframe(table.get('data'))
            
            with st.expander("View extracted text"):
                st.text_area("Text Preview", result['text'][:1000], height=200)
